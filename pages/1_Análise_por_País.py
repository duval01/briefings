import streamlit as st
import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime
import io
import re
import zipfile
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Bloco da Logo na Sidebar ---
logo_sidebar_path = "LogoMinasGerais.png"
if os.path.exists(logo_sidebar_path):
    st.sidebar.image(logo_sidebar_path, width=200)
# --- Fim do Bloco ---

# --- CONFIGURAÇÕES GLOBAIS E CONSTANTES ---
requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)
estados_brasileiros = {'AC', 'AL', 'AP', 'AM', 'BA', 'CE', 'DF', 'ES', 'GO', 'MA', 'MT', 'MS', 'MG', 'PA', 'PB', 'PR',
                      'PE', 'PI', 'RJ', 'RN', 'RS', 'RO', 'RR', 'SC', 'SE', 'SP', 'TO'}
meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}
MESES_MAPA = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
    "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}
LISTA_MESES = list(MESES_MAPA.keys())
ARTIGOS_PAISES_MAP = {
    "Afeganistão": "o", "África do Sul": "a", "Alemanha": "a", "Arábia Saudita": "a",
    "Argentina": "a", "Austrália": "a", "Bélgica": "a", "Brasil": "o", "Canadá": "o",
    "Chade": "o", "Chile": "o", "China": "a", "Colômbia": "a", "Congo": "o",
    "Coreia do Norte": "a", "Coreia do Sul": "a", "Costa Rica": "a", "Equador": "o",
    "Egito": "o", "Emirados Árabes Unidos": "os", "Espanha": "a", "Estados Unidos": "os",
    "Filipinas": "as", "França": "a", "Holanda": "a", "Índia": "a", "Indonésia": "a",
    "Inglaterra": "a", "Irã": "o", "Itália": "a", "Japão": "o", "Líbano": "o",
    "Malásia": "a", "México": "o", "Nicarágua": "a", "Noruega": "a", "Nova Zelândia": "a",
    "Países Baixos": "os", "Panamá": "o", "Paraguai": "o", "Pérsia": "a", "Peru": "o",
    "Reino Unido": "o", "República Checa": "a", "República Dominicana": "a",
    "Romênia": "a", "Rússia": "a", "Singapura": "a", "Suécia": "a", "Uruguai": "o",
    "Venezuela": "a", "Vietnã": "o"
}

# --- BLOCO MANUAL DE BLOCOS ECONÔMICOS ---
# (Como definido anteriormente)
BLOCOS_ECONOMICOS = {
    "América Central e Caribe": [],
    "América do Norte": ["Estados Unidos", "Canadá", "México"],
    "América do Sul": ["Argentina", "Bolívia", "Chile", "Colômbia", "Equador", "Guiana", "Paraguai", "Peru", "Suriname", "Uruguai", "Venezuela"],
    "Associação de Nações do Sudeste Asiático - ASEAN": ["Brunei", "Camboja", "Filipinas", "Indonésia", "Laos", "Malásia", "Myanmar", "Singapura", "Tailândia", "Vietnã"],
    "Comunidade Andina das Nações - CAN": ["Bolívia", "Colômbia", "Equador", "Peru"],
    "Europa": ["Albânia", "Alemanha", "Andorra", "Áustria", "Bélgica", "Bielorrússia", "Bósnia-Herzegovina", "Bulgária", "Chipre", "Croácia", "Dinamarca", "Eslováquia", "Eslovênia", "Espanha", "Estônia", "Finlândia", "França", "Grécia", "Hungria", "Irlanda", "Islândia", "Itália", "Letônia", "Listenstaine", "Lituânia", "Luxemburgo", "Macedônia do Norte", "Malta", "Moldávia", "Mônaco", "Montenegro", "Noruega", "Países Baixos", "Polônia", "Portugal", "Reino Unido", "República Checa", "Romênia", "Rússia", "San Marino", "Sérvia", "Suécia", "Suíça", "Ucrânia", "Vaticano"],
    "Mercado Comum do Sul - Mercosul": ["Argentina", "Paraguai", "Uruguai", "Venezuela"],
    "Oceania": ["Austrália", "Fiji", "Ilhas Marshall", "Ilhas Salomão", "Kiribati", "Micronésia", "Nauru", "Nova Zelândia", "Palau", "Papua Nova Guiné", "Samoa", "Tonga", "Tuvalu", "Vanuatu"],
    "Oriente Médio": ["Arábia Saudita", "Bahrein", "Catar", "Emirados Árabes Unidos", "Iêmen", "Irã", "Iraque", "Israel", "Jordânia", "Kuwait", "Líbano", "Omã", "Palestina", "Síria", "Turquia"],
    "União Europeia - UE": ["Alemanha", "Áustria", "Bélgica", "Bulgária", "Chipre", "Croácia", "Dinamarca", "Eslováquia", "Eslovênia", "Espanha", "Estônia", "Finlândia", "França", "Grécia", "Hungria", "Irlanda", "Itália", "Letônia", "Lituânia", "Luxemburgo", "Malta", "Países Baixos", "Polônia", "Portugal", "República Checa", "Romênia", "Suécia"],
    "África": ["África do Sul", "Angola", "Argélia", "Benin", "Botsuana", "Burkina Faso", "Burundi", "Cabo Verde", "Camarões", "Chade", "Comores", "Congo", "Costa do Marfim", "Djibuti", "Egito", "Eritreia", "Eswatini", "Etiópia", "Gabão", "Gâmbia", "Gana", "Guiné", "Guiné Equatorial", "Guiné-Bissau", "Lesoto", "Libéria", "Líbia", "Madagascar", "Malawi", "Mali", "Marrocos", "Maurício", "Mauritânia", "Moçambique", "Namíbia", "Níger", "Nigéria", "Quênia", "República Centro-Africana", "República Democrática do Congo", "Ruanda", "São Tomé e Príncipe", "Senegal", "Seychelles", "Serra Leoa", "Somália", "Sudão", "Sudão do Sul", "Tanzânia", "Togo", "Tunísia", "Uganda", "Zâmbia", "Zimbábue"],
    "Ásia (Exclusive Oriente Médio)": ["Afeganistão", "Armênia", "Azerbaijão", "Bangladesh", "Brunei", "Butão", "Camboja", "Cazaquistão", "China", "Coreia do Norte", "Coreia do Sul", "Filipinas", "Geórgia", "Índia", "Indonésia", "Japão", "Laos", "Malásia", "Maldivas", "Mongólia", "Myanmar", "Nepal", "Paquistão", "Quirguistão", "Singapura", "Sri Lanka", "Tailândia", "Tajiquistão", "Timor-Leste", "Turcomenistão", "Uzbequistão", "Vietnã"]
}
# --- FIM DO BLOCO MANUAL ---

NCM_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_NCM', 'CO_NCM']
NCM_DTYPES = {'CO_NCM': str, 'CO_SH4': str} 
MUN_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_MUN', 'CO_MUN']
MUN_DTYPES = {'CO_MUN': str}

# --- FUNÇÕES DE LÓGICA (Helpers) ---

@st.cache_data(ttl=3600)
def ler_dados_csv_online(url, usecols=None, dtypes=None):
    retries = 3
    for attempt in range(retries):
        try:
            resposta = requests.get(url, verify=False, timeout=(10, 1200)) 
            resposta.raise_for_status()
            final_dtypes = {'CO_SH4': str, 'CO_NCM': str}
            if dtypes:
                final_dtypes.update(dtypes)
            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1',
                             sep=';', 
                             dtype=final_dtypes,
                             usecols=usecols)
            return df
        except requests.exceptions.RequestException as e:
            print(f"Erro ao acessar o CSV (tentativa {attempt + 1}/{retries}): {e}")
            if "Read timed out" in str(e) and attempt < retries - 1:
                st.warning("Download demorou muito. Tentando novamente...")
                continue
            if "IncompleteRead" in str(e) and attempt < retries - 1:
                st.warning("Retentando download...")
                continue
            else:
                return None
        except Exception as e:
            print(f"Erro inesperado ao baixar ou processar o CSV: {e}")
            return None
    return None

@st.cache_data(ttl=3600)
def carregar_dataframe(url, nome_arquivo, usecols=None, dtypes=None, mostrar_progresso=True):
    progress_bar = None
    if mostrar_progresso: 
        progress_bar = st.progress(0, text=f"Carregando {nome_arquivo}...")
    
    df = ler_dados_csv_online(url, usecols=usecols, dtypes=dtypes)
    
    if mostrar_progresso and progress_bar: 
        if df is not None:
            progress_bar.progress(100, text=f"{nome_arquivo} carregado com sucesso.")
        else:
            progress_bar.empty()
    return df

# --- ALTERADO: Função de Países (padrão do script de Produtos) ---
@st.cache_data
def obter_dados_paises():
    """Carrega a tabela de países (ID e Nome) e armazena em cache."""
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv", usecols=['NO_PAIS', 'CO_PAIS'], mostrar_progresso=False) 
    if df_pais is not None and not df_pais.empty:
        mapa_codigo_nome = pd.Series(df_pais.NO_PAIS.values, index=df_pais.CO_PAIS).to_dict()
        lista_nomes = sorted(df_pais[df_pais['NO_PAIS'] != 'Brasil']['NO_PAIS'].unique().tolist())
        mapa_nome_codigo = pd.Series(df_pais.CO_PAIS.values, index=df_pais.NO_PAIS).to_dict()
        return mapa_codigo_nome, lista_nomes, mapa_nome_codigo
    return {}, [], {}
# --- FIM ALTERADO ---

# --- NOVO: Funções de NCM (padrão do script de Produtos) ---
@st.cache_data
def obter_dados_produtos_ncm():
    """Carrega a tabela NCM completa (SH2, SH4 e SH6) e armazena em cache."""
    url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
    usecols_ncm = ['CO_SH2', 'NO_SH2_POR', 'CO_SH4', 'NO_SH4_POR', 'CO_SH6', 'NO_SH6_POR']
    df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv", usecols=usecols_ncm, mostrar_progresso=False)
    if df_ncm is not None:
        df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].astype(str).str.zfill(4)
        mapa_sh4 = df_ncm.drop_duplicates('CO_SH4_STR').set_index('CO_SH4_STR')['NO_SH4_POR']
        return df_ncm, mapa_sh4.to_dict()
    return None, {}

def get_sh4(co_ncm):
    """Extrai SH4 de um CO_NCM."""
    co_ncm_str = str(co_ncm).strip()
    if pd.isna(co_ncm_str) or co_ncm_str == "":
        return None
    co_ncm_str = co_ncm_str.zfill(8)
    return co_ncm_str[:4]
# --- FIM NOVO ---

@st.cache_data
def obter_lista_de_blocos():
    """Retorna uma lista de nomes de blocos econômicos (hardcoded)."""
    blocos = sorted(list(BLOCOS_ECONOMICOS.keys()))
    return blocos

@st.cache_data
def obter_paises_do_bloco(nome_bloco):
    """Retorna uma lista de nomes de países (hardcoded) para um bloco específico."""
    return BLOCOS_ECONOMICOS.get(nome_bloco, [])

# --- ALTERADO: Função de Países (padrão do script de Produtos) ---
def obter_lista_de_paises(lista_nomes):
    """Apenas retorna a lista de nomes pré-carregada."""
    if not lista_nomes:
        return ["Erro ao carregar lista de países"]
    return lista_nomes
# --- FIM ALTERADO ---

def obter_codigo_pais(nome_pais, mapa_reverso):
    """Obtém o código do país a partir do mapa."""
    return mapa_reverso.get(nome_pais)

# --- ALTERADO: Função de Países (padrão do script de Produtos) ---
def validar_paises(paises_selecionados, mapa_nome_codigo):
    """Valida a lista de países usando o mapa pré-carregado."""
    codigos_paises = []
    nomes_paises_validos = []
    paises_invalidos = []
    
    for pais in paises_selecionados:
        if pais.lower() == "brasil":
            paises_invalidos.append(f"{pais} (Não é possível fazer busca no Brasil)")
            continue
        codigo_pais = mapa_nome_codigo.get(pais) 
        if codigo_pais is None:
            paises_invalidos.append(f"{pais} (País não encontrado)")
        else:
            codigos_paises.append(codigo_pais)
            nomes_paises_validos.append(pais)
    return codigos_paises, nomes_paises_validos, paises_invalidos
# --- FIM ALTERADO ---

def filtrar_dados_por_estado_e_mes(df, estados, meses_para_filtrar):
    df_filtrado = df[df['SG_UF_NCM'].isin(list(estados))]
    df_filtrado = df_filtrado[df_filtrado['CO_MES'].isin(meses_para_filtrar)]
    return df_filtrado

def filtrar_dados_por_mg_e_pais(df, codigos_paises, agrupado, meses_para_filtrar):
    df_filtrado = df[df['SG_UF_NCM'] == 'MG']
    if agrupado:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'].isin(codigos_paises)]
    else:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'] == codigos_paises[0]]
    df_filtrado = df_filtrado[df_filtrado['CO_MES'].isin(meses_para_filtrar)]
    return df_filtrado

def calcular_soma_por_estado(df, df_anterior=None):
    soma_ano = df.groupby('SG_UF_NCM')['VL_FOB'].sum()
    if df_anterior is not None:
        soma_ano_anterior = df_anterior.groupby('SG_UF_NCM')['VL_FOB'].sum()
        return soma_ano, soma_ano_anterior
    else:
        return soma_ano

def calcular_classificacao_estados(soma_ano, soma_ano_anterior, ano_principal, ano_comparacao):
    classificacao = pd.concat([soma_ano, soma_ano_anterior], axis=1, keys=[str(ano_principal), str(ano_comparacao)]).sort_values(
        by=str(ano_principal), ascending=False)
    return classificacao

def calcular_posicao_mg(classificacao):
    posicao_mg = classificacao.index.get_loc('MG') + 1
    return posicao_mg

def calcular_ranking_por_pais(df):
    ranking = df.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
    return ranking

def calcular_participacao(valor_parcial, valor_total):
    if valor_total == 0:
        return 0.0
    participacao = round(valor_parcial / valor_total * 100, 2)
    return participacao

def calcular_diferenca_percentual(valor_atual, valor_anterior):
    if valor_anterior == 0:
        return 0.0, "acréscimo" if valor_atual > 0 else "redução" if valor_atual < 0 else "estabilidade"
    diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    if diferenca > 0:
        tipo_diferenca = "um acréscimo"
    elif diferenca < 0:
        tipo_diferenca = "uma redução"
    else:
        tipo_diferenca = "uma estabilidade"
    diferenca = abs(diferenca)
    return diferenca, tipo_diferenca

def calcular_posicao_estado_pais(df, codigos_paises):
    df_comercio_pais = df[df['CO_PAIS'].isin(codigos_paises)]
    if df_comercio_pais.empty:
        return 0 
    ranking_estados_pais = df_comercio_pais.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(
        ascending=False)
    if 'MG' not in ranking_estados_pais.index:
        return 0 
    posicao_mg_pais = ranking_estados_pais.index.get_loc('MG') + 1
    return posicao_mg_pais

def calcular_balanca_e_fluxo(exportacao_ano, importacao_ano, exportacao_ano_anterior, importacao_ano_anterior):
    balanca_ano = exportacao_ano - importacao_ano
    balanca_ano_anterior = exportacao_ano_anterior - importacao_ano_anterior
    fluxo_comercial_ano = exportacao_ano + importacao_ano
    fluxo_comercial_ano_anterior = exportacao_ano_anterior + importacao_ano_anterior
    variacao_balanca = 0
    variacao_fluxo = 0
    if balanca_ano_anterior != 0:
        variacao_balanca = ((balanca_ano - balanca_ano_anterior) / balanca_ano_anterior) * 100
    if fluxo_comercial_ano_anterior != 0:
        variacao_fluxo = ((fluxo_comercial_ano - fluxo_comercial_ano_anterior) / fluxo_comercial_ano_anterior) * 100
    return balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo

def agregar_dados_por_municipio(df):
    dados_por_municipio = df.groupby('CO_MUN')['VL_FOB'].sum().sort_values(ascending=False)
    total_municipios = dados_por_municipio.sum()
    return dados_por_municipio, total_municipios

def agregar_dados_por_produto(df, df_ncm):
    """Mantida para a geração de texto do DOCX."""
    df_filtered = df.copy()
    
    # Adiciona SH4 se não existir (necessário para esta função)
    if 'SH4' not in df_filtered.columns:
        df_filtered['SH4'] = df_filtered['CO_NCM'].apply(get_sh4).astype(str)
        
    df_sh4_not_null = df_filtered.dropna(subset=['SH4'])
    produtos = df_sh4_not_null.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(5)
    produtos_nomes = {}
    
    # Prepara o df_ncm para busca
    if 'CO_SH4_STR' not in df_ncm.columns:
         df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].astype(str).str.zfill(4)

    for sh4_code, valor in produtos.items():
        filtro_ncm = df_ncm[df_ncm['CO_SH4_STR'] == sh4_code] 
        if not filtro_ncm.empty:
            nome_produto = filtro_ncm['NO_SH4_POR'].iloc[0]
            produtos_nomes[nome_produto] = valor
        else:
            produtos_nomes[f"Produto SH4 {sh4_code} não encontrado"] = valor
    return produtos_nomes

def obter_artigo_pais(nome_pais):
    return ARTIGOS_PAISES_MAP.get(nome_pais) 

def formatar_valor(valor):
    prefixo = ""
    if valor < 0:
        prefixo = "-"
        valor = abs(valor)
    if valor >= 1_000_000_000:
        valor_formatado_str = f"{(valor / 1_000_000_000):.2f}".replace('.',',')
        unidade = "bilhão" if (valor / 1_000_000_000) < 2 else "bilhões"
        return f"{prefixo}US$ {valor_formatado_str} {unidade}"
    if valor >= 1_000_000:
        valor_formatado_str = f"{(valor / 1_000_000):.2f}".replace('.',',')
        unidade = "milhão" if (valor / 1_000_000) < 2 else "milhões"
        return f"{prefixo}US$ {valor_formatado_str} {unidade}"
    if valor >= 1_000:
        valor_formatado_str = f"{(valor / 1_000):.2f}".replace('.',',')
        return f"{prefixo}US$ {valor_formatado_str} mil"
    valor_formatado_str = f"{valor:.2f}".replace('.',',')
    return f"{prefixo}US$ {valor_formatado_str}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

class DocumentoApp:
    def __init__(self, logo_path):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""
        self.logo_path = logo_path
        self.diretorio_base = "/tmp/" 

    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_conteudo_formatado(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_paragrafo(self, texto): 
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_titulo(self, texto):
        p = self.doc.add_paragraph()
        if self.subsecao_atual == 0:
            run = p.add_run(f"{self.secao_atual}. {texto}")
        else:
            run = p.add_run(f"{self.secao_atual}.{self.subsecao_atual}. {texto}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def nova_secao(self):
        self.secao_atual += 1
        self.subsecao_atual = 0

    def nova_subsecao(self):
        self.subsecao_atual += 1

    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        largura_total_cm = 16.0
        table = header.add_table(rows=1, cols=2, width=Cm(largura_total_cm))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(4.0)
        table.columns[1].width = Cm(12.0)
        cell_imagem = table.cell(0, 0)
        paragraph_imagem = cell_imagem.paragraphs[0]
        paragraph_imagem.paragraph_format.space_before = Pt(0)
        paragraph_imagem.paragraph_format.space_after = Pt(0)
        run_imagem = paragraph_imagem.add_run()
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                run_imagem.add_picture(self.logo_path,
                                       width=Cm(3.5), 
                                       height=Cm(3.42))
            except Exception as e:
                paragraph_imagem.add_run("[Logo não encontrado]")
        else:
            paragraph_imagem.add_run("[Logo não encontrado]")
        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        def formatar_paragrafo_cabecalho(p):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p = cell_texto.paragraphs[0]
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[0])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True 
        p = cell_texto.add_paragraph()
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[1])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        for texto in textos[2:]: 
            p = cell_texto.add_paragraph()
            formatar_paragrafo_cabecalho(p)
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = False 

    def finalizar_documento(self):
        diretorio_real = self.diretorio_base
        try:
            os.makedirs(diretorio_real, exist_ok=True)
        except Exception:
            diretorio_real = "/tmp/"
            os.makedirs(diretorio_real, exist_ok=True)
        nome_arquivo = f"{self.titulo_doc}.docx"
        nome_arquivo_sanitizado = sanitize_filename(nome_arquivo)
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        file_bytes = file_stream.getvalue()
        st.success(f"Documento '{nome_arquivo_sanitizado}' gerado com sucesso!")
        try:
            self.doc.save(os.path.join(diretorio_real, nome_arquivo_sanitizado))
        except Exception:
            pass 
        return file_bytes, nome_arquivo_sanitizado

# --- ----------------------------------- ---
# --- INTERFACE GRÁFICA DO STREAMLIT (Página 1) ---
# --- ----------------------------------- ---

# --- Inicialização do Session State ---
if 'arquivos_gerados_pais' not in st.session_state:
    st.session_state.arquivos_gerados_pais = []

# --- Callback para limpar o state ---
def clear_download_state_pais():
    if 'arquivos_gerados_pais' in st.session_state:
        st.session_state.arquivos_gerados_pais = []

# --- ENTRADAS PRINCIPAIS ---
st.header("1. Configurações da Análise")


# --- ALTERADO: Carrega mapas de nomes ---
try:
    mapa_nomes_paises, lista_paises_nomes, mapa_paises_reverso = obter_dados_paises()
    lista_de_blocos = obter_lista_de_blocos()
    _, mapa_sh4_nomes = obter_dados_produtos_ncm() # Para as tabelas
except Exception as e:
    st.error(f"Erro crítico ao carregar listas iniciais: {e}")
    lista_paises_nomes = ["Falha ao carregar países"]
    lista_de_blocos = ["Falha ao carregar blocos"]
    mapa_nomes_paises = {}
    mapa_paises_reverso = {}
    mapa_sh4_nomes = {}
# --- FIM ALTERADO ---

lista_de_paises = obter_lista_de_paises(lista_paises_nomes)

# --- Lógica de 'default' resiliente ---
valores_padrao = ["China", "Estados Unidos"]
valores_padrao_filtrados = [pais for pais in valores_padrao if pais in lista_de_paises]
if not valores_padrao_filtrados and len(lista_de_paises) > 0 and "Erro" not in lista_de_paises[0]:
    valores_padrao_filtrados = [lista_de_paises[0]]
elif "Erro" in lista_de_paises[0] or "Falha" in lista_de_paises[0]:
    valores_padrao_filtrados = [] 
    st.warning("Não foi possível carregar la lista de países. O site de dados pode estar fora do ar.")
# --- FIM DA LÓGICA ---

ano_atual = datetime.now().year

col1, col2 = st.columns(2)
with col1:
    ano_principal = st.number_input(
        "Ano de Referência:",
        min_value=1998,
        max_value=ano_atual,
        value=ano_atual,
        help="O ano principal que você quer analisar.",
        on_change=clear_download_state_pais
    )
    ano_comparacao = st.number_input(
        "Ano de Comparação:",
        min_value=1998,
        max_value=ano_atual,
        value=ano_atual - 1,
        help="O ano contra o qual você quer comparar.",
        on_change=clear_download_state_pais
    )
    meses_selecionados = st.multiselect(
        "Meses de Análise (opcional):",
        options=LISTA_MESES,
        help="Selecione os meses. Se deixar em branco, o ano inteiro será analisado.",
        on_change=clear_download_state_pais
    )
    
    # --- NOVO: Seletor de N Produtos ---
    top_n_produtos = st.number_input(
        "Nº de Produtos no Ranking:",
        min_value=1,
        max_value=100,
        value=10,      # Valor padrão
        help="Quantos produtos (SH4) devem ser exibidos nas tabelas de ranking (Top 10, Top 20, etc.).",
        on_change=clear_download_state_pais
    )
    # --- FIM NOVO ---

with col2:
    blocos_selecionados = st.multiselect(
        "Filtrar por Bloco(s) (opcional):",
        options=lista_de_blocos,
        help="Os países destes blocos serão adicionados à seleção.",
        on_change=clear_download_state_pais
    )
    
    paises_selecionados_manual = st.multiselect(
        "Filtrar por País(es) (opcional):",
        options=lista_de_paises,
        default=valores_padrao_filtrados,
        help="Você pode digitar para pesquisar e selecionar múltiplos países.",
        on_change=clear_download_state_pais
    )

# --- LÓGICA CONDICIONAL PARA ENTRADAS ---
agrupado = True 
nome_agrupamento = None

paises_do_bloco = []
if blocos_selecionados:
    for bloco in blocos_selecionados:
        paises_do_bloco.extend(obter_paises_do_bloco(bloco))

paises = sorted(list(set(paises_selecionados_manual + paises_do_bloco)))

if len(paises) > 1:
    st.header("2. Opções de Agrupamento")
    
    if blocos_selecionados and not paises_selecionados_manual:
        agrupado = True
        st.info(f"Análise de Bloco será agrupada.")
        nome_agrupamento = ", ".join(blocos_selecionados)
    else:
        agrupamento_input = st.radio(
            "Deseja que os dados sejam agrupados ou separados?",
            ("agrupados", "separados"),
            index=0,
            horizontal=True,
            on_change=clear_download_state_pais
        )
        agrupado = (agrupamento_input == "agrupados")

        if agrupado:
            # --- INSERÇÃO DA DICA (PAÍSES) ---
            st.info(
                "💡 **Como funciona o agrupamento:**\n"
                "* **Agrupados:** Gerará um **único relatório** consolidado. As tabelas de ranking de produtos mostrarão a **soma** de todos os países selecionados. O detalhamento por país aparecerá no expansor.\n"
                "* **Separados:** Gerará um **relatório individual** para cada país. O download será um arquivo .zip."
            )
            # --- FIM DA INSERÇÃO ---
            
            quer_nome_agrupamento = st.checkbox(
                "Deseja dar um nome para este agrupamento?", 
                key="pais_nome_grupo",
                on_change=clear_download_state_pais
            )
            if quer_nome_agrupamento:
                nome_agrupamento = st.text_input(
                    "Digite o nome do agrupamento:", 
                    key="pais_nome_input",
                    on_change=clear_download_state_pais
                )
    st.header("3. Gerar Relatório")
else:
    agrupado = False 
    st.header("2. Gerar Relatório")


# --- EXECUÇÃO DO SCRIPT ---
if st.button(" Iniciar Geração do Relatório"):
    
    st.session_state.arquivos_gerados_pais = []
    
    logo_path_to_use = "LogoMinasGerais.png" 
    if not os.path.exists(logo_path_to_use):
        st.warning(f"Aviso: A logo 'LogoMinasGerais.png' não foi encontrada. O cabeçalho será gerado sem a logo.")
        logo_path_to_use = None
    
    with st.spinner(f"Gerando relatório para {', '.join(paises)} ({ano_principal} vs {ano_comparacao})... Isso pode levar alguns minutos."):
        
        try:
            # --- ALTERADO: Passa o mapa reverso para validação ---
            codigos_paises, nomes_paises_validos, paises_invalidos = validar_paises(paises, mapa_paises_reverso)
            if paises_invalidos:
                st.warning(f"Países não encontrados ou inválidos (ignorados): {', '.join(paises_invalidos)}")
            if not nomes_paises_validos:
                st.error("Nenhum país válido fornecido. A geração foi interrompida.")
                st.stop()
            
            url_exp_ano_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_principal}.csv"
            url_exp_ano_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_comparacao}.csv"
            url_imp_ano_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_principal}.csv"
            url_imp_ano_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_comparacao}.csv"
            url_ncm_sh = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
            url_exp_mun_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_principal}_MUN.csv"
            url_imp_mun_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_principal}_MUN.csv"
            url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
            
            # --- ALTERADO: Carrega df_ncm principal aqui ---
            df_ncm, _ = carregar_dataframe(url_ncm_sh, "NCM_SH.csv", usecols=['CO_SH4', 'NO_SH4_POR'], dtypes={'CO_SH4': str}, mostrar_progresso=False)
            df_uf_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['CO_MUN_GEO', 'NO_MUN_MIN'], mostrar_progresso=False)
            
            if df_ncm is None or df_uf_mun is None:
                st.error("Não foi possível carregar tabelas auxiliares (NCM ou Municípios). Abortando.")
                st.stop()
            # Prepara o df_ncm para busca
            df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].astype(str).str.zfill(4)
            # --- FIM ALTERADO ---

            df_exp_ano = carregar_dataframe(url_exp_ano_principal, f"EXP_{ano_principal}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            df_exp_ano_anterior = carregar_dataframe(url_exp_ano_comparacao, f"EXP_{ano_comparacao}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)

            if df_exp_ano is None or df_exp_ano_anterior is None:
                st.error("Não foi possível carregar dados de exportação. Verifique os anos selecionados ou tente novamente mais tarde.")
                st.stop()
                
            # --- NOVO: Adiciona colunas SH4 ---
            df_exp_ano['SH4'] = df_exp_ano['CO_NCM'].apply(get_sh4)
            df_exp_ano_anterior['SH4'] = df_exp_ano_anterior['CO_NCM'].apply(get_sh4)
            # --- FIM NOVO ---

            ultimo_mes_disponivel = df_exp_ano['CO_MES'].max()
            meses_para_filtrar = []
            
            if not meses_selecionados: 
                meses_para_filtrar = list(range(1, ultimo_mes_disponivel + 1))
                nome_periodo = f"o ano de {ano_principal} (até {meses_pt[ultimo_mes_disponivel]})"
                nome_periodo_em = f"No ano de {ano_principal} (até {meses_pt[ultimo_mes_disponivel]})"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            else:
                meses_para_filtrar = [MESES_MAPA[m] for m in meses_selecionados]
                if max(meses_para_filtrar) > ultimo_mes_disponivel:
                    st.error(f"O ano {ano_principal} só possui dados até {meses_pt[ultimo_mes_disponivel]}. Por favor, desmarque os meses posteriores.")
                    st.stop()
                nome_periodo = f"o período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_em = f"No período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            
            df_exp_ano_estados = filtrar_dados_por_estado_e_mes(df_exp_ano, estados_brasileiros, meses_para_filtrar)
            df_exp_ano_anterior_estados = filtrar_dados_por_estado_e_mes(df_exp_ano_anterior, estados_brasileiros, meses_para_filtrar)
            df_exp_ano_mg = filtrar_dados_por_estado_e_mes(df_exp_ano, ['MG'], meses_para_filtrar)
            df_exp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano, codigos_paises, agrupado, meses_para_filtrar)
            df_exp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior, codigos_paises, agrupado, meses_para_filtrar)
            exportacao_pais_ano = df_exp_ano_mg_paises['VL_FOB'].sum()
            exportacao_pais_ano_anterior = df_exp_ano_anterior_mg_paises['VL_FOB'].sum()
            
            if not agrupado:
                ranking_mg_exp = calcular_ranking_por_pais(df_exp_ano_mg)
                try:
                    posicao_pais_mg_exp = ranking_mg_exp.index.get_loc(codigos_paises[0]) + 1
                except KeyError:
                    posicao_pais_mg_exp = 0
            else:
                posicao_pais_mg_exp = None
            
            exportacao_mg_total_ano = df_exp_ano_mg['VL_FOB'].sum()
            participacao_pais_mg_exp = calcular_participacao(exportacao_pais_ano, exportacao_mg_total_ano)
            diferenca_exportacao, tipo_diferenca_exp = calcular_diferenca_percentual(exportacao_pais_ano, exportacao_pais_ano_anterior)
            
            exportacao_mg_para_pais = df_exp_ano[ (df_exp_ano['SG_UF_NCM'] == 'MG') & (df_exp_ano['CO_PAIS'].isin(codigos_paises)) & (df_exp_ano['CO_MES'].isin(meses_para_filtrar)) ]['VL_FOB'].sum()
            exportacao_brasil_pais = df_exp_ano[ (df_exp_ano['CO_PAIS'].isin(codigos_paises)) & (df_exp_ano['CO_MES'].isin(meses_para_filtrar)) ]['VL_FOB'].sum()
            
            participacao_mg_brasil_exp = calcular_participacao(exportacao_mg_para_pais, exportacao_brasil_pais)
            posicao_mg_pais_exp = calcular_posicao_estado_pais(df_exp_ano_estados, codigos_paises)
            produtos_exportacao = agregar_dados_por_produto(df_exp_ano_mg_paises.copy(), df_ncm) # Para o texto
            
            df_exp_mun = carregar_dataframe(url_exp_mun_principal, f"EXP_{ano_principal}_MUN.csv", usecols=MUN_COLS, dtypes=MUN_DTYPES)
            if df_exp_mun is None:
                st.error("Não foi possível carregar dados de exportação por município. Abortando.")
                st.stop()
                
            df_exp_mun_filtrado = df_exp_mun[(df_exp_mun['SG_UF_MUN'] == 'MG') & (df_exp_mun['CO_PAIS'].isin(codigos_paises)) & (df_exp_mun['CO_MES'].isin(meses_para_filtrar))]
            exportacoes_por_municipio, total_exportacoes_municipios = agregar_dados_por_municipio(df_exp_mun_filtrado)
            
            df_imp_ano = carregar_dataframe(url_imp_ano_principal, f"IMP_{ano_principal}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            df_imp_ano_anterior = carregar_dataframe(url_imp_ano_comparacao, f"IMP_{ano_comparacao}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            
            if df_imp_ano is None or df_imp_ano_anterior is None:
                st.error("Não foi possível carregar dados de importação. Abortando.")
                st.stop()
            
            # --- NOVO: Adiciona colunas SH4 ---
            df_imp_ano['SH4'] = df_imp_ano['CO_NCM'].apply(get_sh4)
            df_imp_ano_anterior['SH4'] = df_imp_ano_anterior['CO_NCM'].apply(get_sh4)
            # --- FIM NOVO ---
            
            df_imp_ano_estados = filtrar_dados_por_estado_e_mes(df_imp_ano, estados_brasileiros, meses_para_filtrar)
            df_imp_ano_anterior_estados = filtrar_dados_por_estado_e_mes(df_imp_ano_anterior, estados_brasileiros, meses_para_filtrar)
            df_imp_ano_mg = filtrar_dados_por_estado_e_mes(df_imp_ano, ['MG'], meses_para_filtrar)
            df_imp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano, codigos_paises, agrupado, meses_para_filtrar)
            df_imp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior, codigos_paises, agrupado, meses_para_filtrar)
            importacao_pais_ano = df_imp_ano_mg_paises['VL_FOB'].sum()
            importacao_pais_ano_anterior = df_imp_ano_anterior_mg_paises['VL_FOB'].sum()
            
            if not agrupado:
                ranking_mg_imp = calcular_ranking_por_pais(df_imp_ano_mg)
                try:
                    posicao_pais_mg_imp = ranking_mg_imp.index.get_loc(codigos_paises[0]) + 1
                except KeyError:
                    posicao_pais_mg_imp = 0
            else:
                posicao_pais_mg_imp = None
            
            importacao_mg_total_ano = df_imp_ano_mg['VL_FOB'].sum()
            participacao_pais_mg_imp = calcular_participacao(importacao_pais_ano, importacao_mg_total_ano)
            diferenca_importacao, tipo_diferenca_imp = calcular_diferenca_percentual(importacao_pais_ano, importacao_pais_ano_anterior)
            
            importacao_mg_para_pais = df_imp_ano[ (df_imp_ano['SG_UF_NCM'] == 'MG') & (df_imp_ano['CO_PAIS'].isin(codigos_paises)) & (df_imp_ano['CO_MES'].isin(meses_para_filtrar)) ]['VL_FOB'].sum()
            importacao_brasil_pais = df_imp_ano[ (df_imp_ano['CO_PAIS'].isin(codigos_paises)) & (df_imp_ano['CO_MES'].isin(meses_para_filtrar)) ]['VL_FOB'].sum()
            
            participacao_mg_brasil_imp = calcular_participacao(importacao_mg_para_pais, importacao_brasil_pais)
            posicao_mg_pais_imp = calcular_posicao_estado_pais(df_imp_ano_estados, codigos_paises)
            produtos_importacao = agregar_dados_por_produto(df_imp_ano_mg_paises.copy(), df_ncm) # Para o texto
            
            df_imp_mun = carregar_dataframe(url_imp_mun_principal, f"IMP_{ano_principal}_MUN.csv", usecols=MUN_COLS, dtypes=MUN_DTYPES)
            if df_imp_mun is None:
                st.error("Não foi possível carregar dados de importação por município. Abortando.")
                st.stop()

            df_imp_mun_filtrado = df_imp_mun[(df_imp_mun['SG_UF_MUN'] == 'MG') & (df_imp_mun['CO_PAIS'].isin(codigos_paises)) & (df_imp_mun['CO_MES'].isin(meses_para_filtrar))]
            importacoes_por_municipio, total_importacoes_municipios = agregar_dados_por_municipio(df_imp_mun_filtrado)
            
            balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo = calcular_balanca_e_fluxo(exportacao_pais_ano, importacao_pais_ano, exportacao_pais_ano_anterior, importacao_pais_ano_anterior)
            
            if agrupado:
                app = DocumentoApp(logo_path=logo_path_to_use)
                paises_corretos = nomes_paises_validos 
                nome_relatorio = nome_agrupamento if (nome_agrupamento and nome_agrupamento.strip() != "") else ', '.join(paises_corretos)

                # Geração de Texto
                fluxo_e_balanca = f"Considerando {nome_periodo}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, ..."
                texto_exportacao = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)}..."
                # ... (resto da geração de texto) ...
                
                # Montagem do Documento
                titulo_documento = f"Briefing - {nome_relatorio} - {ano_principal}"
                app.set_titulo(titulo_documento)
                # ... (resto da montagem do docx) ...
                
                # --- NOVO: Lógica das Tabelas (Agrupado) ---
                
                # --- Tabela Exportação ---
                st.header("Principais Produtos Exportados (MG para Agrupamento)")
                exp_produtos_princ = df_exp_ano_mg_paises.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).reset_index()
                exp_produtos_comp = df_exp_ano_anterior_mg_paises.groupby('SH4')['VL_FOB'].sum().reset_index()
                
                exp_produtos_princ['Produto'] = exp_produtos_princ['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                exp_produtos_princ[f'Valor {ano_principal} (US$)'] = exp_produtos_princ['VL_FOB']
                exp_produtos_comp['Produto'] = exp_produtos_comp['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                exp_produtos_comp[f'Valor {ano_comparacao} (US$)'] = exp_produtos_comp['VL_FOB']
                
                exp_final_prod = pd.merge(exp_produtos_princ[['Produto', 'SH4', f'Valor {ano_principal} (US$)']], 
                                          exp_produtos_comp[['Produto', 'SH4', f'Valor {ano_comparacao} (US$)']], 
                                          on=["Produto", "SH4"], how="outer").fillna(0)
                
                exp_final_prod['Variação %'] = 100 * (exp_final_prod[f'Valor {ano_principal} (US$)'] - exp_final_prod[f'Valor {ano_comparacao} (US$)']) / exp_final_prod[f'Valor {ano_comparacao} (US$)']
                exp_final_prod['Variação %'] = exp_final_prod['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                exp_final_prod[f'Valor {ano_principal}'] = exp_final_prod[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                exp_final_prod[f'Valor {ano_comparacao}'] = exp_final_prod[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                
                df_display_exp_prod = exp_final_prod.sort_values(by=f'Valor {ano_principal} (US$)', ascending=False).reset_index(drop=True)
                
                st.dataframe(
                    df_display_exp_prod[['Produto', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_produtos),
                    hide_index=True
                )

                # --- Expander Exportação ---
                with st.expander("Ver detalhamento de países por produto (Exportação)"):
                    top_produtos_lista = df_display_exp_prod['SH4'].head(top_n_produtos).tolist()
                    
                    detalhe_exp_princ = df_exp_ano_mg_paises.groupby(['SH4', 'CO_PAIS'])['VL_FOB'].sum().reset_index()
                    detalhe_exp_comp = df_exp_ano_anterior_mg_paises.groupby(['SH4', 'CO_PAIS'])['VL_FOB'].sum().reset_index()
                    
                    detalhe_exp_princ['Produto'] = detalhe_exp_princ['SH4'].map(mapa_sh4_nomes)
                    detalhe_exp_comp['Produto'] = detalhe_exp_comp['SH4'].map(mapa_sh4_nomes)
                    detalhe_exp_princ['País'] = detalhe_exp_princ['CO_PAIS'].map(mapa_nomes_paises)
                    detalhe_exp_comp['País'] = detalhe_exp_comp['CO_PAIS'].map(mapa_nomes_paises)

                    detalhe_exp_princ = detalhe_exp_princ.rename(columns={'VL_FOB': f'Valor {ano_principal} (US$)'})
                    detalhe_exp_comp = detalhe_exp_comp.rename(columns={'VL_FOB': f'Valor {ano_comparacao} (US$)'})

                    detalhe_exp_final = pd.merge(
                        detalhe_exp_princ[['Produto', 'País', 'SH4', f'Valor {ano_principal} (US$)']],
                        detalhe_exp_comp[['Produto', 'País', 'SH4', f'Valor {ano_comparacao} (US$)']],
                        on=['Produto', 'País', 'SH4'], how='outer'
                    ).fillna(0)

                    detalhe_exp_final = detalhe_exp_final[detalhe_exp_final['SH4'].isin(top_produtos_lista)]

                    detalhe_exp_final['Variação %'] = 100 * (detalhe_exp_final[f'Valor {ano_principal} (US$)'] - detalhe_exp_final[f'Valor {ano_comparacao} (US$)']) / detalhe_exp_final[f'Valor {ano_comparacao} (US$)']
                    detalhe_exp_final['Variação %'] = detalhe_exp_final['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                    detalhe_exp_final[f'Valor {ano_principal}'] = detalhe_exp_final[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                    detalhe_exp_final[f'Valor {ano_comparacao}'] = detalhe_exp_final[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)

                    detalhe_exp_final = detalhe_exp_final.sort_values(by=['Produto', f'Valor {ano_principal} (US$)'], ascending=[True, False])
                    st.dataframe(
                        detalhe_exp_final[['Produto', 'País', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']],
                        hide_index=True
                    )
                
                # --- Tabela Importação ---
                st.header("Principais Produtos Importados (MG do Agrupamento)")
                imp_produtos_princ = df_imp_ano_mg_paises.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).reset_index()
                imp_produtos_comp = df_imp_ano_anterior_mg_paises.groupby('SH4')['VL_FOB'].sum().reset_index()
                
                imp_produtos_princ['Produto'] = imp_produtos_princ['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                imp_produtos_princ[f'Valor {ano_principal} (US$)'] = imp_produtos_princ['VL_FOB']
                imp_produtos_comp['Produto'] = imp_produtos_comp['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                imp_produtos_comp[f'Valor {ano_comparacao} (US$)'] = imp_produtos_comp['VL_FOB']
                
                imp_final_prod = pd.merge(imp_produtos_princ[['Produto', 'SH4', f'Valor {ano_principal} (US$)']], 
                                          imp_produtos_comp[['Produto', 'SH4', f'Valor {ano_comparacao} (US$)']], 
                                          on=["Produto", "SH4"], how="outer").fillna(0)
                
                imp_final_prod['Variação %'] = 100 * (imp_final_prod[f'Valor {ano_principal} (US$)'] - imp_final_prod[f'Valor {ano_comparacao} (US$)']) / imp_final_prod[f'Valor {ano_comparacao} (US$)']
                imp_final_prod['Variação %'] = imp_final_prod['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                imp_final_prod[f'Valor {ano_principal}'] = imp_final_prod[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                imp_final_prod[f'Valor {ano_comparacao}'] = imp_final_prod[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                
                df_display_imp_prod = imp_final_prod.sort_values(by=f'Valor {ano_principal} (US$)', ascending=False).reset_index(drop=True)
                
                st.dataframe(
                    df_display_imp_prod[['Produto', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_produtos),
                    hide_index=True
                )

                # --- Expander Importação ---
                with st.expander("Ver detalhamento de países por produto (Importação)"):
                    top_produtos_lista_imp = df_display_imp_prod['SH4'].head(top_n_produtos).tolist()
                    
                    detalhe_imp_princ = df_imp_ano_mg_paises.groupby(['SH4', 'CO_PAIS'])['VL_FOB'].sum().reset_index()
                    detalhe_imp_comp = df_imp_ano_anterior_mg_paises.groupby(['SH4', 'CO_PAIS'])['VL_FOB'].sum().reset_index()
                    
                    detalhe_imp_princ['Produto'] = detalhe_imp_princ['SH4'].map(mapa_sh4_nomes)
                    detalhe_imp_comp['Produto'] = detalhe_imp_comp['SH4'].map(mapa_sh4_nomes)
                    detalhe_imp_princ['País'] = detalhe_imp_princ['CO_PAIS'].map(mapa_nomes_paises)
                    detalhe_imp_comp['País'] = detalhe_imp_comp['CO_PAIS'].map(mapa_nomes_paises)

                    detalhe_imp_princ = detalhe_imp_princ.rename(columns={'VL_FOB': f'Valor {ano_principal} (US$)'})
                    detalhe_imp_comp = detalhe_imp_comp.rename(columns={'VL_FOB': f'Valor {ano_comparacao} (US$)'})

                    detalhe_imp_final = pd.merge(
                        detalhe_imp_princ[['Produto', 'País', 'SH4', f'Valor {ano_principal} (US$)']],
                        detalhe_imp_comp[['Produto', 'País', 'SH4', f'Valor {ano_comparacao} (US$)']],
                        on=['Produto', 'País', 'SH4'], how='outer'
                    ).fillna(0)

                    detalhe_imp_final = detalhe_imp_final[detalhe_imp_final['SH4'].isin(top_produtos_lista_imp)]

                    detalhe_imp_final['Variação %'] = 100 * (detalhe_imp_final[f'Valor {ano_principal} (US$)'] - detalhe_imp_final[f'Valor {ano_comparacao} (US$)']) / detalhe_imp_final[f'Valor {ano_comparacao} (US$)']
                    detalhe_imp_final['Variação %'] = detalhe_imp_final['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                    detalhe_imp_final[f'Valor {ano_principal}'] = detalhe_imp_final[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                    detalhe_imp_final[f'Valor {ano_comparacao}'] = detalhe_imp_final[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)

                    detalhe_imp_final = detalhe_imp_final.sort_values(by=['Produto', f'Valor {ano_principal} (US$)'], ascending=[True, False])
                    st.dataframe(
                        detalhe_imp_final[['Produto', 'País', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']],
                        hide_index=True
                    )
                # --- FIM NOVO ---
                
                # Gera o Docx (código original omitido para brevidade)
                app.set_titulo(titulo_documento)
                # ...
                file_bytes, file_name = app.finalizar_documento() 
                st.session_state.arquivos_gerados_pais.append({"name": file_name, "data": file_bytes})

            else:
                # --- LÓGICA PARA SEPARADOS ---
                paises_corretos = nomes_paises_validos
                
                for pais in paises_corretos:
                    st.subheader(f"Processando: {pais}") 
                    app = DocumentoApp(logo_path=logo_path_to_use)
                    
                    codigos_paises_loop = [obter_codigo_pais(pais, mapa_paises_reverso)] # Usa o mapa

                    # Recarrega os dados filtrados para este país
                    df_exp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano, codigos_paises_loop, False, meses_para_filtrar)
                    df_exp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior, codigos_paises_loop, False, meses_para_filtrar)
                    df_imp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano, codigos_paises_loop, False, meses_para_filtrar)
                    df_imp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior, codigos_paises_loop, False, meses_para_filtrar)
                    
                    # Gera texto (código original omitido para brevidade)
                    # ...
                    
                    # --- NOVO: Lógica das Tabelas (Separados) ---
                    
                    # --- Tabela Exportação ---
                    st.header(f"Principais Produtos Exportados (MG para {pais})")
                    exp_produtos_princ = df_exp_ano_mg_paises.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).reset_index()
                    exp_produtos_comp = df_exp_ano_anterior_mg_paises.groupby('SH4')['VL_FOB'].sum().reset_index()
                    
                    exp_produtos_princ['Produto'] = exp_produtos_princ['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                    exp_produtos_princ[f'Valor {ano_principal} (US$)'] = exp_produtos_princ['VL_FOB']
                    exp_produtos_comp['Produto'] = exp_produtos_comp['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                    exp_produtos_comp[f'Valor {ano_comparacao} (US$)'] = exp_produtos_comp['VL_FOB']
                    
                    exp_final_prod = pd.merge(exp_produtos_princ[['Produto', f'Valor {ano_principal} (US$)']], 
                                              exp_produtos_comp[['Produto', f'Valor {ano_comparacao} (US$)']], 
                                              on="Produto", how="outer").fillna(0)
                    
                    exp_final_prod['Variação %'] = 100 * (exp_final_prod[f'Valor {ano_principal} (US$)'] - exp_final_prod[f'Valor {ano_comparacao} (US$)']) / exp_final_prod[f'Valor {ano_comparacao} (US$)']
                    exp_final_prod['Variação %'] = exp_final_prod['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                    exp_final_prod[f'Valor {ano_principal}'] = exp_final_prod[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                    exp_final_prod[f'Valor {ano_comparacao}'] = exp_final_prod[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                    
                    df_display_exp_prod = exp_final_prod.sort_values(by=f'Valor {ano_principal} (US$)', ascending=False).reset_index(drop=True)
                    
                    st.dataframe(
                        df_display_exp_prod[['Produto', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_produtos),
                        hide_index=True
                    )
                    
                    # --- Tabela Importação ---
                    st.header(f"Principais Produtos Importados (MG de {pais})")
                    imp_produtos_princ = df_imp_ano_mg_paises.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).reset_index()
                    imp_produtos_comp = df_imp_ano_anterior_mg_paises.groupby('SH4')['VL_FOB'].sum().reset_index()
                    
                    imp_produtos_princ['Produto'] = imp_produtos_princ['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                    imp_produtos_princ[f'Valor {ano_principal} (US$)'] = imp_produtos_princ['VL_FOB']
                    imp_produtos_comp['Produto'] = imp_produtos_comp['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                    imp_produtos_comp[f'Valor {ano_comparacao} (US$)'] = imp_produtos_comp['VL_FOB']
                    
                    imp_final_prod = pd.merge(imp_produtos_princ[['Produto', f'Valor {ano_principal} (US$)']], 
                                              imp_produtos_comp[['Produto', f'Valor {ano_comparacao} (US$)']], 
                                              on="Produto", how="outer").fillna(0)
                    
                    imp_final_prod['Variação %'] = 100 * (imp_final_prod[f'Valor {ano_principal} (US$)'] - imp_final_prod[f'Valor {ano_comparacao} (US$)']) / imp_final_prod[f'Valor {ano_comparacao} (US$)']
                    imp_final_prod['Variação %'] = imp_final_prod['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                    imp_final_prod[f'Valor {ano_principal}'] = imp_final_prod[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                    imp_final_prod[f'Valor {ano_comparacao}'] = imp_final_prod[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                    
                    df_display_imp_prod = imp_final_prod.sort_values(by=f'Valor {ano_principal} (US$)', ascending=False).reset_index(drop=True)
                    
                    st.dataframe(
                        df_display_imp_prod[['Produto', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_produtos),
                        hide_index=True
                    )
                    # --- FIM NOVO ---
                    
                    # Gera o Docx (código original omitido para brevidade)
                    # ...
                    file_bytes, file_name = app.finalizar_documento()
                    st.session_state.arquivos_gerados_pais.append({"name": file_name, "data": file_bytes})
            
            # Limpa DFs grandes da memória
            del df_exp_ano, df_exp_ano_anterior, df_imp_ano, df_imp_ano_anterior
                
        except Exception as e:
            st.error(f"Ocorreu um erro inesperado durante a geração:")
            st.exception(e)

# --- Bloco de exibição de Download (COM LÓGICA DE ZIP) ---
if st.session_state.arquivos_gerados_pais:
    st.header("4. Relatórios Gerados")
    st.info("Clique para baixar os relatórios. Eles permanecerão aqui até que você gere um novo relatório.")
    
    if len(st.session_state.arquivos_gerados_pais) > 1:
        st.subheader("Pacote de Relatórios (ZIP)")
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for arquivo in st.session_state.arquivos_gerados_pais:
                zip_file.writestr(arquivo["name"], arquivo["data"])
        
        zip_bytes = zip_buffer.getvalue()
        
        st.download_button(
            label=f"Baixar todos os {len(st.session_state.arquivos_gerados_pais)} relatórios (.zip)",
            data=zip_bytes,
            file_name=f"Briefings_Países_{ano_principal}.zip",
            mime="application/zip",
            key="download_zip_pais"
        )
        
    elif len(st.session_state.arquivos_gerados_pais) == 1:
        st.subheader("Relatório Gerado")
        arquivo = st.session_state.arquivos_gerados_pais[0] 
        st.download_button(
            label=f"Baixar Relatório ({arquivo['name']})",
            data=arquivo["data"], 
            file_name=arquivo["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{arquivo['name']}"
        )

# --- Bloco de Rodapé ---
st.divider() 

col1, col2 = st.columns([0.3, 0.7], vertical_alignment="center") 

with col1:
    logo_footer_path = "AEST Sede.png"
    if os.path.exists(logo_footer_path):
        st.image(logo_footer_path, width=150)
    else:
        st.caption("Logo AEST não encontrada.")

with col2:
    st.caption("Desenvolvido por Aest - Dados e Subsecretaria de Promoção de Investimentos e Cadeias Produtivas")
