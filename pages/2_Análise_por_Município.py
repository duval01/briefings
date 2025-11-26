import streamlit as st
import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime
import io
import re
import zipfile
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- OCULTA A NAVEGAÇÃO PADRÃO ---
st.markdown(
    """
    <style>
        [data-testid="stSidebarNav"] {
            display: none;
        }
    </style>
    """,
    unsafe_allow_html=True
)

# --- IMPORTAÇÃO E PROTEÇÃO DA PÁGINA ---
try:
    from Home import draw_sidebar, logout
except ImportError:
    def draw_sidebar():
        st.sidebar.error("Erro ao carregar a navegação. Execute a partir do Home.py.")
    def logout():
        st.sidebar.error("Erro ao carregar.")

st.session_state.current_page = 'Análise por Município'
draw_sidebar()

if not st.session_state.get('logged_in', False):
    st.error("Acesso negado. Por favor, faça o login na Página Principal.")
    st.page_link("Home.py", label="Ir para a página de Login", icon="🏠")
    st.stop()
# --- FIM DA PROTEÇÃO ---

# --- CONFIGURAÇÕES GLOBAIS ---
requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)

MESES_MAPA = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
    "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}
LISTA_MESES = list(MESES_MAPA.keys())
meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# --- DICIONÁRIO DE MESORREGIÕES DE MG ---
MESORREGIOES_MG = {
    "Noroeste de Minas": ["Unaí", "Paracatu", "João Pinheiro", "Vazante", "Brasilândia de Minas", "Guarda-Mor", "Arinos", "Buritis"],
    "Norte de Minas": ["Montes Claros", "Janaúba", "Januária", "Pirapora", "Salinas", "Bocaiúva", "Porteirinha", "Jaíba", "Várzea da Palma", "Grão Mogol"],
    "Jequitinhonha": ["Diamantina", "Almenara", "Capelinha", "Araçuaí", "Minas Novas", "Itamarandiba", "Pedra Azul", "Jequitinhonha"],
    "Vale do Mucuri": ["Teófilo Otoni", "Nanuque", "Carlos Chagas", "Malacacheta", "Águas Formosas", "Itaipé"],
    "Triângulo Mineiro e Alto Paranaíba": ["Uberlândia", "Uberaba", "Araxá", "Patos de Minas", "Ituiutaba", "Frutal", "Patrocínio", "Araguari", "Conceição das Alagoas", "Sacramento", "Iturama"],
    "Central Mineira": ["Curvelo", "Três Marias", "Bom Despacho", "Felixlândia", "Corinto", "Pompeu", "Morada Nova de Minas"],
    "Metropolitana de Belo Horizonte": ["Belo Horizonte", "Betim", "Contagem", "Sete Lagoas", "Nova Lima", "Santa Luzia", "Ribeirão das Neves", "Ibirité", "Sabará", "Vespasiano", "Itabira", "Ouro Preto", "Mariana", "Congonhas", "Conselheiro Lafaiete"],
    "Vale do Rio Doce": ["Governador Valadares", "Ipatinga", "Coronel Fabriciano", "Timóteo", "Caratinga", "Aimorés", "Mantena", "Resplendor"],
    "Oeste de Minas": ["Divinópolis", "Formiga", "Itaúna", "Pará de Minas", "Nova Serrana", "Arcos", "Bambuí", "Piumhi", "Campo Belo", "Oliveira"],
    "Sul e Sudoeste de Minas": ["Poços de Caldas", "Pouso Alegre", "Varginha", "Passos", "Itajubá", "Alfenas", "Três Corações", "Lavras", "São Sebastião do Paraíso", "Guaxupé", "Extrema", "Varginha"],
    "Campos das Vertentes": ["Barbacena", "São João del Rei", "Lavras", "São Tiago", "Nazareno", "Barroso", "Resende Costa"],
    "Zona da Mata": ["Juiz de Fora", "Ubá", "Muriaé", "Manhuaçu", "Viçosa", "Cataguases", "Ponte Nova", "Leopoldina", "Santos Dumont", "Além Paraíba"]
}

def obter_lista_de_mesorregioes():
    return sorted(list(MESORREGIOES_MG.keys()))

def obter_municipios_da_meso(nome_meso):
    return MESORREGIOES_MG.get(nome_meso, [])

# --- DEFINIÇÃO DE COLUNAS E TIPOS ---
MUN_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_MUN', 'CO_MUN', 'SH4', 'CO_SH4', 'CO_NCM'] 
MUN_DTYPES = {'CO_MUN': str, 'CO_SH4': str, 'CO_NCM': str, 'SH4': str, 'CO_PAIS': str}

# --- FUNÇÕES DE LÓGICA (Helpers) ---

# --- FUNÇÃO REINSERIDA: Normaliza códigos ---
def normalizar_codigo(codigo):
    """Remove .0, espaços e converte para string limpa."""
    if pd.isna(codigo) or codigo == "": return None
    s = str(codigo).strip()
    if s.endswith('.0'):
        s = s[:-2]
    return s
# --------------------------------------------

@st.cache_data(ttl=3600)
def ler_dados_csv_online(url, usecols=None, dtypes=None):
    retries = 3
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    for attempt in range(retries):
        try:
            resposta = requests.get(url, headers=headers, verify=False, timeout=(10, 1200)) 
            resposta.raise_for_status()
            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1',
                             sep=';', dtype=dtypes)
            
            if 'CO_ANO' not in df.columns and '<!DOCTYPE' in str(df.columns):
                return None 
            
            if usecols:
                cols_existentes = [c for c in usecols if c in df.columns]
                df = df[cols_existentes]
                
            return df
        except (requests.exceptions.RequestException, ConnectionResetError) as e:
            print(f"Erro ao acessar o CSV (tentativa {attempt + 1}/{retries}): {e}")
            if attempt < retries - 1:
                import time
                time.sleep(2)
                continue
            else:
                return None
        except Exception as e:
            print(f"Erro inesperado ao baixar ou processar o CSV: {e}")
            return None
    return None

@st.cache_data(ttl=3600)
def carregar_dataframe(url, nome_arquivo, usecols=None, dtypes=None, mostrar_progresso=True):
    progress_bar = None
    if mostrar_progresso: 
        progress_bar = st.progress(0, text=f"Carregando {nome_arquivo}...")
    
    df = ler_dados_csv_online(url, usecols=None, dtypes=dtypes)
    
    if mostrar_progresso and progress_bar: 
        if df is not None:
            progress_bar.progress(100, text=f"{nome_arquivo} carregado com sucesso.")
        else:
            progress_bar.empty()

    return df

@st.cache_data
def obter_dados_paises():
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv", usecols=['NO_PAIS', 'CO_PAIS'], dtypes={'CO_PAIS': str}, mostrar_progresso=False) 
    if df_pais is not None and not df_pais.empty:
        # Usa normalizar_codigo aqui
        df_pais['CO_PAIS'] = df_pais['CO_PAIS'].apply(normalizar_codigo)
        return pd.Series(df_pais.NO_PAIS.values, index=df_pais.CO_PAIS).to_dict()
    return {}

@st.cache_data
def obter_lista_de_municipios():
    url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
    df_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['SG_UF', 'NO_MUN', 'CO_MUN_GEO'], mostrar_progresso=False)
    if df_mun is not None:
        lista_mun = df_mun[df_mun['SG_UF'] == 'MG']['NO_MUN'].unique().tolist()
        lista_mun.sort()
        return lista_mun
    return ["Erro ao carregar lista de municípios"]

@st.cache_data
def obter_mapa_codigos_municipios():
    url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
    df_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", 
                                usecols=['SG_UF', 'NO_MUN', 'CO_MUN_GEO'], 
                                dtypes={'CO_MUN_GEO': str}, 
                                mostrar_progresso=False)
    if df_mun is not None:
        df_mun_mg = df_mun[df_mun['SG_UF'] == 'MG'].copy()
        df_mun_mg['CO_MUN_GEO'] = df_mun_mg['CO_MUN_GEO'].apply(normalizar_codigo)
        return pd.Series(df_mun_mg.CO_MUN_GEO.values, index=df_mun_mg.NO_MUN).to_dict()
    return {}

@st.cache_data
def obter_dados_produtos_ncm():
    url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
    usecols_ncm = ['CO_SH2', 'NO_SH2_POR', 'CO_SH4', 'NO_SH4_POR']
    df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv", usecols=usecols_ncm, dtypes={'CO_SH4': str, 'CO_SH2': str}, mostrar_progresso=False)
    
    if df_ncm is not None:
        df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].apply(normalizar_codigo).str.zfill(4)
        df_ncm['CO_SH2_STR'] = df_ncm['CO_SH2'].apply(normalizar_codigo).str.zfill(2)
        
        mapa_sh4 = df_ncm.drop_duplicates('CO_SH4_STR').set_index('CO_SH4_STR')['NO_SH4_POR'].to_dict()
        mapa_sh2 = df_ncm.drop_duplicates('CO_SH2_STR').set_index('CO_SH2_STR')['NO_SH2_POR'].to_dict()
        
        return mapa_sh4, mapa_sh2
    return {}, {}

def get_sh4(codigo):
    s = normalizar_codigo(codigo)
    if not s: return None
    return s.zfill(4)[:4]

def get_sh2(sh4):
    s = normalizar_codigo(sh4)
    if s: return s[:2]
    return None

def normalizar_coluna_produto(df):
    if df is None: return None
    for col in ['CO_MUN', 'CO_PAIS', 'CO_SH4', 'CO_NCM']:
        if col in df.columns:
            df[col] = df[col].apply(normalizar_codigo)

    if 'SH4' in df.columns: 
        df['SH4'] = df['SH4'].apply(get_sh4)
        return df
        
    if 'CO_SH4' in df.columns:
        df['SH4'] = df['CO_SH4'].apply(get_sh4)
    elif 'CO_NCM' in df.columns:
        df['SH4'] = df['CO_NCM'].apply(get_sh4)
    else:
        df['SH4'] = '0000'
    return df

def formatar_valor(valor):
    if pd.isna(valor): return "US$ 0,00"
    prefixo = ""
    if valor < 0:
        prefixo = "-"
        valor = abs(valor)
    if valor >= 1_000_000_000:
        return f"{prefixo}US$ {(valor / 1_000_000_000):.2f} bilhões"
    if valor >= 1_000_000:
        return f"{prefixo}US$ {(valor / 1_000_000):.2f} milhões"
    if valor >= 1_000:
        return f"{prefixo}US$ {(valor / 1_000):.2f} mil"
    return f"{prefixo}US$ {valor:.2f}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def calc_var_display(row, col_atual, col_ant):
    v_atual = row[col_atual]
    v_ant = row[col_ant]
    if pd.isna(v_ant) or v_ant == 0:
        return "Novo Mercado" if v_atual > 0 else "-"
    var = ((v_atual - v_ant) / v_ant) * 100
    return f"{var:.2f}%"

def calcular_diferenca_percentual(valor_atual, valor_anterior):
    if pd.isna(valor_anterior) or valor_anterior == 0:
        return 100.0 if valor_atual > 0 else 0.0, "acréscimo" if valor_atual > 0 else "estabilidade"
        
    diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    tipo_diferenca = "um acréscimo" if diferenca > 0 else "uma redução" if diferenca < 0 else "uma estabilidade"
    return abs(diferenca), tipo_diferenca

class DocumentoApp:
    def __init__(self, logo_path):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""
        self.logo_path = logo_path
        self.diretorio_base = "/tmp/" 

    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_conteudo_formatado(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_paragrafo(self, texto): 
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_titulo(self, texto):
        p = self.doc.add_paragraph()
        if self.subsecao_atual == 0:
            run = p.add_run(f"{self.secao_atual}. {texto}")
        else:
            run = p.add_run(f"{self.secao_atual}.{self.subsecao_atual}. {texto}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def nova_secao(self):
        self.secao_atual += 1
        self.subsecao_atual = 0

    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Cm(16.0))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(4.0)
        table.columns[1].width = Cm(12.0)
        cell_imagem = table.cell(0, 0)
        paragraph_imagem = cell_imagem.paragraphs[0]
        run_imagem = paragraph_imagem.add_run()
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                run_imagem.add_picture(self.logo_path, width=Cm(3.5), height=Cm(3.42))
            except: pass
        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        def formatar_paragrafo_cabecalho(p):
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, texto in enumerate(textos):
            p = cell_texto.paragraphs[0] if i == 0 else cell_texto.add_paragraph()
            formatar_paragrafo_cabecalho(p)
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = (i < 2)

    def finalizar_documento(self):
        try: os.makedirs(self.diretorio_base, exist_ok=True)
        except: pass
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue(), f"{sanitize_filename(self.titulo_doc)}.docx"

# --- CONFIGURAÇÃO DA PÁGINA ---
st.header("1. Configurações da Análise Municipal")

def clear_download_state_mun():
    if 'arquivos_gerados_municipio' in st.session_state:
        st.session_state.arquivos_gerados_municipio = []

# Carregamento de dados auxiliares
lista_de_municipios = obter_lista_de_municipios()
mapa_codigos_municipios = obter_mapa_codigos_municipios()
mapa_nomes_paises = obter_dados_paises()
mapa_sh4_nomes, mapa_sh2_nomes = obter_dados_produtos_ncm()
lista_de_mesorregioes = obter_lista_de_mesorregioes()
ano_atual = datetime.now().year

# --- LAYOUT BALANCEADO (3 x 3) ---
col1, col2 = st.columns(2)
with col1:
    ano_principal = st.number_input(
        "Ano de Referência:", min_value=1998, max_value=ano_atual, value=ano_atual,
        on_change=clear_download_state_mun
    )
    mesorregioes_selecionadas = st.multiselect(
        "Filtrar por Mesorregião (opcional):",
        options=lista_de_mesorregioes,
        on_change=clear_download_state_mun
    )
    municipios_selecionados = st.multiselect(
        "Selecione o(s) município(s):",
        options=lista_de_municipios,
        default=["BELO HORIZONTE"],
        help="Você pode digitar para pesquisar.",
        on_change=clear_download_state_mun
    )

with col2:
    ano_comparacao = st.number_input(
        "Ano de Comparação:", min_value=1998, max_value=ano_atual, value=ano_atual - 1,
        on_change=clear_download_state_mun
    )
    meses_selecionados = st.multiselect(
        "Meses de Análise (opcional):",
        options=LISTA_MESES,
        on_change=clear_download_state_mun
    )
    top_n_itens = st.number_input(
        "Nº de Itens nos Rankings:",
        min_value=1,
        max_value=100,
        value=10,
        help="Número de países/produtos a exibir nas tabelas.",
        on_change=clear_download_state_mun
    )
# --- FIM DO LAYOUT ---

# --- Lógica de Agrupamento ---
if mesorregioes_selecionadas:
    municipios_da_meso = []
    for meso in mesorregioes_selecionadas:
        municipios_da_meso.extend(obter_municipios_da_meso(meso))
    todos_municipios = list(set(municipios_selecionados + municipios_da_meso))
else:
    todos_municipios = municipios_selecionados

agrupado = True
nome_agrupamento = None
if mesorregioes_selecionadas and not municipios_selecionados:
    nome_sugerido = ", ".join(mesorregioes_selecionadas)
else:
    nome_sugerido = ""

if len(todos_municipios) > 1:
    st.header("2. Opções de Agrupamento")
    agrupamento_input = st.radio(
        f"Deseja que os dados dos {len(todos_municipios)} municípios sejam agrupados?",
        ("agrupados", "separados"),
        index=0,
        horizontal=True,
        on_change=clear_download_state_mun
    )
    agrupado = (agrupamento_input == "agrupados")
    
    if agrupado:
        st.info("Agrupados: Relatório consolidado. Separados: Relatórios individuais.")
        quer_nome_agrupamento = st.checkbox("Deseja dar um nome para este agrupamento?", value=bool(nome_sugerido))
        if quer_nome_agrupamento:
            nome_agrupamento = st.text_input("Digite o nome do agrupamento:", value=nome_sugerido)
    st.header("3. Gerar Análise")
else:
    agrupado = False 
    st.header("2. Gerar Análise")


# --- EXECUÇÃO ---
if st.button("Iniciar Análise por Município"):
    st.session_state.arquivos_gerados_municipio = []
    logo_path_to_use = "LogoMinasGerais.png"
    
    with st.spinner(f"Processando {len(todos_municipios)} municípios..."):
        try:
            # --- Validação ---
            codigos_municipios_map = []
            municipios_validos = []
            for m in todos_municipios:
                cod = mapa_codigos_municipios.get(m) or mapa_codigos_municipios.get(m.upper()) or mapa_codigos_municipios.get(m.title())
                if cod:
                    codigos_municipios_map.append(normalizar_codigo(cod))
                    municipios_validos.append(m)
            
            if not codigos_municipios_map:
                st.error("Nenhum município válido.")
                st.stop()

            # --- URLs e Carregamento ---
            url_exp_mun_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_principal}_MUN.csv"
            url_exp_mun_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_comparacao}_MUN.csv"
            url_imp_mun_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_principal}_MUN.csv"
            url_imp_mun_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_comparacao}_MUN.csv"

            df_exp_mun_princ = carregar_dataframe(url_exp_mun_principal, f"EXP_{ano_principal}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)
            df_exp_mun_comp = carregar_dataframe(url_exp_mun_comparacao, f"EXP_{ano_comparacao}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)
            df_imp_mun_princ = carregar_dataframe(url_imp_mun_principal, f"IMP_{ano_principal}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)
            df_imp_mun_comp = carregar_dataframe(url_imp_mun_comparacao, f"IMP_{ano_comparacao}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)

            if df_exp_mun_princ is None:
                st.error("Falha ao carregar dados.")
                st.stop()

            # --- Normalização ---
            df_exp_mun_princ = normalizar_coluna_produto(df_exp_mun_princ)
            df_exp_mun_comp = normalizar_coluna_produto(df_exp_mun_comp)
            df_imp_mun_princ = normalizar_coluna_produto(df_imp_mun_princ)
            df_imp_mun_comp = normalizar_coluna_produto(df_imp_mun_comp)

            for df in [df_exp_mun_princ, df_exp_mun_comp, df_imp_mun_princ, df_imp_mun_comp]:
                if df is not None:
                    df['SH2'] = df['SH4'].astype(str).str[:2]
                    if 'CO_MUN' in df.columns:
                        df['CO_MUN'] = df['CO_MUN'].apply(normalizar_codigo)

            # --- Filtro Mês ---
            if meses_selecionados:
                meses_para_filtrar = [MESES_MAPA[m] for m in meses_selecionados]
                nome_periodo = f"o período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            else:
                meses_para_filtrar = list(range(1, df_exp_mun_princ['CO_MES'].max() + 1))
                nome_periodo = f"o ano de {ano_principal} (completo)"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"

            # --- Loop ---
            if not agrupado:
                municipios_para_processar = municipios_validos
            else:
                municipios_para_processar = [nome_agrupamento if (nome_agrupamento and nome_agrupamento.strip() != "") else ", ".join(municipios_validos)]

            for municipio_nome in municipios_para_processar:
                app = DocumentoApp(logo_path=logo_path_to_use)
                
                if agrupado:
                    st.subheader(f"Análise Agrupada: {municipio_nome}")
                    codigos_loop = codigos_municipios_map
                    nome_limpo = sanitize_filename(municipio_nome)
                    titulo_doc = f"Briefing - {nome_limpo} - {ano_principal}"
                    nome_doc = f"de {municipio_nome}"
                else:
                    st.subheader(f"Análise: {municipio_nome}")
                    c = mapa_codigos_municipios.get(municipio_nome) or mapa_codigos_municipios.get(municipio_nome.upper())
                    codigos_loop = [normalizar_codigo(c)]
                    nome_limpo = sanitize_filename(municipio_nome)
                    titulo_doc = f"Briefing - {nome_limpo} - {ano_principal}"
                    nome_doc = f"de {municipio_nome}"
                
                app.set_titulo(titulo_doc)

                # --- EXPORTAÇÃO ---
                st.header(f"Exportações ({municipio_nome})")
                df_exp_princ_f = df_exp_mun_princ[(df_exp_mun_princ['CO_MUN'].isin(codigos_loop)) & (df_exp_mun_princ['CO_MES'].isin(meses_para_filtrar))]
                df_exp_comp_f = df_exp_mun_comp[(df_exp_mun_comp['CO_MUN'].isin(codigos_loop)) & (df_exp_mun_comp['CO_MES'].isin(meses_para_filtrar))]
                
                exp_total_princ = df_exp_princ_f['VL_FOB'].sum()
                exp_total_comp = df_exp_comp_f['VL_FOB'].sum()
                dif_exp, tipo_dif_exp = calcular_diferenca_percentual(exp_total_princ, exp_total_comp)

                # Tabela Países
                st.subheader("Principais Destinos")
                exp_paises = df_exp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                exp_paises_c = df_exp_comp_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                
                exp_paises['País'] = exp_paises['CO_PAIS'].apply(normalizar_codigo).map(mapa_nomes_paises).fillna("Desconhecido")
                exp_paises_c['País'] = exp_paises_c['CO_PAIS'].apply(normalizar_codigo).map(mapa_nomes_paises).fillna("Desconhecido")
                
                exp_final = pd.merge(exp_paises, exp_paises_c, on='País', how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                col_princ = f'VL_FOB {ano_principal}'
                col_comp = f'VL_FOB {ano_comparacao}'
                
                exp_final['Variação %'] = exp_final.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                exp_final = exp_final.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}'}).sort_values(by=f'Valor {ano_principal}', ascending=False)
                
                df_show_exp = exp_final.copy()
                df_show_exp[f'Valor {ano_principal}'] = df_show_exp[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_exp[f'Valor {ano_comparacao}'] = df_show_exp[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_exp.head(top_n_itens), hide_index=True, use_container_width=True)

                # Tabela Produtos
                st.subheader("Principais Produtos Exportados")
                exp_prod = df_exp_princ_f.groupby(['SH4', 'SH2'])['VL_FOB'].sum().reset_index()
                exp_prod_c = df_exp_comp_f.groupby(['SH4', 'SH2'])['VL_FOB'].sum().reset_index()
                
                exp_prod['Descrição SH4'] = exp_prod['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                exp_prod['Descrição SH2'] = exp_prod['SH2'].map(mapa_sh2_nomes).fillna("Desconhecido")
                
                exp_final_p = pd.merge(exp_prod, exp_prod_c, on=['SH4', 'SH2'], how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                # Preenche descrições perdidas no merge
                exp_final_p['Descrição SH4'] = exp_final_p['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                exp_final_p['Descrição SH2'] = exp_final_p['SH2'].map(mapa_sh2_nomes).fillna("Desconhecido")

                exp_final_p['Variação %'] = exp_final_p.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                exp_final_p = exp_final_p.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}', 'SH4': 'Código SH4', 'SH2': 'Código SH2'})
                exp_final_p = exp_final_p.sort_values(by=f'Valor {ano_principal}', ascending=False)

                df_show_p = exp_final_p.copy()
                df_show_p[f'Valor {ano_principal}'] = df_show_p[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_p[f'Valor {ano_comparacao}'] = df_show_p[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_p[['Código SH4', 'Descrição SH4', 'Código SH2', 'Descrição SH2', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_itens), hide_index=True, use_container_width=True)
                
                # Expander
                with st.expander("Ver detalhamento (País + Produto)"):
                    detalhe = df_exp_princ_f.groupby(['CO_PAIS', 'SH4'])['VL_FOB'].sum().reset_index()
                    # CORREÇÃO AQUI: Usa normalizar_codigo no map
                    detalhe['País'] = detalhe['CO_PAIS'].apply(normalizar_codigo).map(mapa_nomes_paises)
                    detalhe['Produto'] = detalhe['SH4'].map(mapa_sh4_nomes)
                    detalhe['Valor'] = detalhe['VL_FOB'].apply(formatar_valor)
                    st.dataframe(detalhe.sort_values('VL_FOB', ascending=False), use_container_width=True)

                # Docx
                texto_exp = f"Em {nome_periodo}, as exportações {nome_doc} somaram {formatar_valor(exp_total_princ)}, {tipo_dif_exp} de {dif_exp:.1f}% em relação a {nome_periodo_comp}."
                app.nova_secao()
                app.adicionar_titulo("Exportações")
                app.adicionar_conteudo_formatado(texto_exp)

                # === IMPORTAÇÃO ===
                st.header(f"Importações ({municipio_nome})")
                
                df_imp_princ_f = df_imp_mun_princ[(df_imp_mun_princ['CO_MUN'].isin(codigos_loop)) & (df_imp_mun_princ['CO_MES'].isin(meses_para_filtrar))]
                df_imp_comp_f = df_imp_mun_comp[(df_imp_mun_comp['CO_MUN'].isin(codigos_loop)) & (df_imp_mun_comp['CO_MES'].isin(meses_para_filtrar))]
                
                imp_total_princ = df_imp_princ_f['VL_FOB'].sum()
                imp_total_comp = df_imp_comp_f['VL_FOB'].sum()
                dif_imp, tipo_dif_imp = calcular_diferenca_percentual(imp_total_princ, imp_total_comp)

                # Tabela Países
                st.subheader("Principais Origens")
                imp_paises = df_imp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                imp_paises_c = df_imp_comp_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                
                imp_paises['País'] = imp_paises['CO_PAIS'].apply(normalizar_codigo).map(mapa_nomes_paises).fillna("Desconhecido")
                imp_paises_c['País'] = imp_paises_c['CO_PAIS'].apply(normalizar_codigo).map(mapa_nomes_paises).fillna("Desconhecido")
                
                imp_final = pd.merge(imp_paises, imp_paises_c, on='País', how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                imp_final['Variação %'] = imp_final.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                imp_final = imp_final.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}'}).sort_values(by=f'Valor {ano_principal}', ascending=False)
                
                df_show_imp = imp_final.copy()
                df_show_imp[f'Valor {ano_principal}'] = df_show_imp[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_imp[f'Valor {ano_comparacao}'] = df_show_imp[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_imp.head(top_n_itens), hide_index=True, use_container_width=True)

                # Tabela Produtos
                st.subheader("Principais Produtos Importados")
                imp_prod = df_imp_princ_f.groupby(['SH4', 'SH2'])['VL_FOB'].sum().reset_index()
                imp_prod_c = df_imp_comp_f.groupby(['SH4', 'SH2'])['VL_FOB'].sum().reset_index()
                
                imp_final_pi = pd.merge(imp_prod, imp_prod_c, on=['SH4', 'SH2'], how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                imp_final_pi['Descrição SH4'] = imp_final_pi['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                imp_final_pi['Descrição SH2'] = imp_final_pi['SH2'].map(mapa_sh2_nomes).fillna("Desconhecido")
                
                imp_final_pi['Variação %'] = imp_final_pi.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                imp_final_pi = imp_final_pi.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}', 'SH4': 'Código SH4', 'SH2': 'Código SH2'}).sort_values(by=f'Valor {ano_principal}', ascending=False)

                df_show_pi = imp_final_pi.copy()
                df_show_pi[f'Valor {ano_principal}'] = df_show_pi[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_pi[f'Valor {ano_comparacao}'] = df_show_pi[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_pi[['Código SH4', 'Descrição SH4', 'Código SH2', 'Descrição SH2', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_itens), hide_index=True, use_container_width=True)

                # Expander
                with st.expander("Ver detalhamento (País + Produto)"):
                    detalhe_i = df_imp_princ_f.groupby(['CO_PAIS', 'SH4'])['VL_FOB'].sum().reset_index()
                    detalhe_i['País'] = detalhe_i['CO_PAIS'].apply(normalizar_codigo).map(mapa_nomes_paises)
                    detalhe_i['Produto'] = detalhe_i['SH4'].map(mapa_sh4_nomes)
                    detalhe_i['Valor'] = detalhe_i['VL_FOB'].apply(formatar_valor)
                    st.dataframe(detalhe_i.sort_values('VL_FOB', ascending=False), use_container_width=True)

                # Docx Importação
                texto_imp = f"Em {nome_periodo}, as importações {nome_doc} somaram {formatar_valor(imp_total_princ)}, {tipo_dif_imp} de {dif_imp:.1f}% em relação a {nome_periodo_comp}."
                app.nova_secao()
                app.adicionar_titulo("Importações")
                app.adicionar_conteudo_formatado(texto_imp)

                # Salvar
                file_bytes, file_name = app.finalizar_documento()
                st.session_state.arquivos_gerados_municipio.append({"name": file_name, "data": file_bytes})

        except Exception as e:
            st.error("Ocorreu um erro.")
            st.exception(e)

# --- Bloco de Download (com ZIP) ---
if st.session_state.arquivos_gerados_municipio:
    st.header("4. Relatórios Gerados")
    st.info("Clique para baixar os relatórios. Eles permanecerão aqui até que você gere um novo relatório.")
    
    if len(st.session_state.arquivos_gerados_municipio) > 1:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for arquivo in st.session_state.arquivos_gerados_municipio:
                zip_file.writestr(arquivo["name"], arquivo["data"])
        
        st.download_button(
            label=f"Baixar todos os {len(st.session_state.arquivos_gerados_municipio)} relatórios (.zip)",
            data=zip_buffer.getvalue(),
            file_name=f"Briefings_Municipios_{ano_principal}.zip",
            mime="application/zip",
            key="download_zip_municipio"
        )
        
    elif len(st.session_state.arquivos_gerados_municipio) == 1:
        arquivo = st.session_state.arquivos_gerados_municipio[0] 
        st.download_button(
            label=f"Baixar Relatório ({arquivo['name']})",
            data=arquivo["data"], 
            file_name=arquivo["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{arquivo['name']}"
        )

# --- Bloco de Rodapé ---
st.divider() 

col1, col2 = st.columns([0.3, 0.7], vertical_alignment="center") 

with col1:
    logo_footer_path = "AEST Sede.png"
    if os.path.exists(logo_footer_path):
        st.image(logo_footer_path, width=150)
    else:
        st.caption("Logo AEST não encontrada.")

with col2:
    st.caption("Desenvolvido por Aest - Dados e Subsecretaria de Promoção de Investimentos e Cadeias Produtivas")