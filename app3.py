import streamlit as st
import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime, timezone
import calendar
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import json
from google.api_core.client_options import ClientOptions
from googleapiclient import discovery
import io
import zipfile

# --- CONFIGURAÇÕES GLOBAIS E CONSTANTES ---

requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)

estados_brasileiros = {'AC', 'AL', 'AP', 'AM', 'BA', 'CE', 'DF', 'ES', 'GO', 'MA', 'MT', 'MS', 'MG', 'PA', 'PB', 'PR',
                      'PE', 'PI', 'RJ', 'RN', 'RS', 'RO', 'RR', 'SC', 'SE', 'SP', 'TO'}
meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# --- LISTAS DE COLUNAS PARA OTIMIZAÇÃO ---
NCM_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_NCM', 'CO_NCM']
NCM_DTYPES = {'CO_NCM': str, 'CO_SH4': str} 

MUN_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_MUN', 'CO_MUN']
MUN_DTYPES = {'CO_MUN': str}


# --- FUNÇÕES DE LÓGICA (OTIMIZADAS) ---

@st.cache_data(ttl=3600)
def ler_dados_csv_online(url, usecols=None, dtypes=None):
    """Lê dados CSV da URL com retentativas e colunas/dtypes específicos."""
    retries = 3
    for attempt in range(retries):
        try:
            resposta = requests.get(url, verify=False, timeout=(10, 1200)) 
            resposta.raise_for_status()
            
            final_dtypes = {'CO_SH4': str, 'CO_NCM': str}
            if dtypes:
                final_dtypes.update(dtypes)

            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1',
                             sep=';', 
                             dtype=final_dtypes,
                             usecols=usecols)
            return df
        except requests.exceptions.RequestException as e:
            st.error(f"Erro ao acessar o CSV (tentativa {attempt + 1}/{retries}): {e}")
            if "Read timed out" in str(e) and attempt < retries - 1:
                st.warning("Download demorou muito. Tentando novamente...")
                continue
            if "IncompleteRead" in str(e) and attempt < retries - 1:
                st.warning("Retentando download...")
                continue
            else:
                st.error(f"Falha ao baixar após {retries} tentativas.")
                return None
        except Exception as e:
            st.error(f"Erro inesperado ao baixar ou processar o CSV: {e}")
            return None
    return None

@st.cache_data(ttl=3600)
def carregar_dataframe(url, nome_arquivo, usecols=None, dtypes=None):
    """Carrega o DataFrame da URL (usa cache) com colunas e dtypes."""
    progress_bar = st.progress(0, text=f"Carregando {nome_arquivo}...")
    df = ler_dados_csv_online(url, usecols=usecols, dtypes=dtypes)
    if df is not None:
        progress_bar.progress(100, text=f"{nome_arquivo} carregado com sucesso.")
    else:
        progress_bar.empty()
    return df

@st.cache_data
def obter_codigo_pais(nome_pais):
    """Obtém o código do país a partir do nome."""
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv", usecols=['NO_PAIS', 'CO_PAIS']) 
    if df_pais is not None and not df_pais.empty:
        filtro_pais = df_pais[df_pais['NO_PAIS'] == nome_pais]
        if not filtro_pais.empty:
            codigo_pais = filtro_pais['CO_PAIS'].iloc[0]
            return codigo_pais
    return None

def validar_paises(paises):
    """Valida a lista de países e retorna códigos e nomes válidos."""
    codigos_paises = []
    nomes_paises_validos = []
    paises_invalidos = []

    for pais in paises:
        if pais.lower() == "brasil":
            paises_invalidos.append(f"{pais} (Não é possível fazer busca no Brasil)")
            continue
            
        codigo_pais = obter_codigo_pais(pais)
        if codigo_pais is None:
            paises_invalidos.append(f"{pais} (País não encontrado)")
        else:
            codigos_paises.append(codigo_pais)
            nomes_paises_validos.append(pais)

    return codigos_paises, nomes_paises_validos, paises_invalidos

def filtrar_dados_por_estado_e_mes(df, estados, ultimo_mes_disponivel, ano_completo):
    """Filtra o DataFrame por estado e, se necessário, por mês."""
    df_filtrado = df[df['SG_UF_NCM'].isin(list(estados))]
    if not ano_completo:
        df_filtrado = df_filtrado[df_filtrado['CO_MES'] <= ultimo_mes_disponivel]
    return df_filtrado

def filtrar_dados_por_mg_e_pais(df, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo):
    """Filtra o DataFrame por MG e pelos países (se agrupado) ou por país individual (se não agrupado) e por mês se ano incompleto."""
    df_filtrado = df[df['SG_UF_NCM'] == 'MG']
    if agrupado:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'].isin(codigos_paises)]
    else:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'] == codigos_paises[0]]
    if not ano_completo:
        df_filtrado = df_filtrado[df_filtrado['CO_MES'] <= ultimo_mes_disponivel]
    return df_filtrado

def calcular_soma_por_estado(df, ano, ano_completo, ultimo_mes_disponivel, df_anterior=None):
    """Calcula a soma dos valores por estado para o ano e, opcionalmente, para o ano anterior."""
    soma_ano = df.groupby('SG_UF_NCM')['VL_FOB'].sum()
    if df_anterior is not None:
        df_anterior_filtered = df_anterior
        if not ano_completo:
            df_anterior_filtered = df_anterior[df_anterior['CO_MES'] <= ultimo_mes_disponivel]
        soma_ano_anterior = df_anterior_filtered.groupby('SG_UF_NCM')['VL_FOB'].sum()
        return soma_ano, soma_ano_anterior
    else:
        return soma_ano

def calcular_classificacao_estados(soma_ano, soma_ano_anterior, ano):
    """Calcula a classificação dos estados por valor total."""
    classificacao = pd.concat([soma_ano, soma_ano_anterior], axis=1, keys=[str(ano), str(ano - 1)]).sort_values(
        by=str(ano), ascending=False)
    return classificacao

def calcular_posicao_mg(classificacao):
    """Calcula a posição de MG na classificação."""
    posicao_mg = classificacao.index.get_loc('MG') + 1
    return posicao_mg

def calcular_ranking_por_pais(df, ano_completo, ultimo_mes_disponivel):
    """Calcula a posição do país no ranking de MG."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]
    ranking = df_filtered.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
    return ranking

def calcular_participacao(valor_parcial, valor_total):
    """Calcula a participação percentual."""
    if valor_total == 0:
        return 0.0
    participacao = round(valor_parcial / valor_total * 100, 2)
    return participacao

def calcular_diferenca_percentual(valor_atual, valor_anterior, ano_completo, ultimo_mes_disponivel):
    """Calcula a diferença percentual entre dois valores."""
    if valor_anterior == 0:
        return 0.0, "acréscimo" if valor_atual > 0 else "redução" if valor_atual < 0 else "estabilidade"
    if ano_completo:
        diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    else:
        diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    if diferenca > 0:
        tipo_diferenca = "um acréscimo"
    elif diferenca < 0:
        tipo_diferenca = "uma redução"
    else:
        tipo_diferenca = "uma estabilidade"
    diferenca = abs(diferenca)
    return diferenca, tipo_diferenca

def calcular_posicao_estado_pais(df, codigos_paises, ano_completo, ultimo_mes_disponivel):
    """Calcula a posição de MG nas exportações/importações para o(s) país(es) informado(s)."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]

    df_comercio_pais = df_filtered[df_filtered['CO_PAIS'].isin(codigos_paises)]

    if df_comercio_pais.empty:
        return 0 

    ranking_estados_pais = df_comercio_pais.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(
        ascending=False)

    if 'MG' not in ranking_estados_pais.index:
        return 0 

    posicao_mg_pais = ranking_estados_pais.index.get_loc('MG') + 1
    return posicao_mg_pais

def calcular_balanca_e_fluxo(exportacao_ano, importacao_ano, exportacao_ano_anterior, importacao_ano_anterior):
    """Calcula a balança comercial, o fluxo comercial e suas variações."""
    balanca_ano = exportacao_ano - importacao_ano
    balanca_ano_anterior = exportacao_ano_anterior - importacao_ano_anterior
    fluxo_comercial_ano = exportacao_ano + importacao_ano
    fluxo_comercial_ano_anterior = exportacao_ano_anterior + importacao_ano_anterior
    variacao_balanca = 0
    variacao_fluxo = 0
    if balanca_ano_anterior != 0:
        variacao_balanca = ((balanca_ano - balanca_ano_anterior) / balanca_ano_anterior) * 100
    if fluxo_comercial_ano_anterior != 0:
        variacao_fluxo = ((fluxo_comercial_ano - fluxo_comercial_ano_anterior) / fluxo_comercial_ano_anterior) * 100

    return balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo

def agregar_dados_por_municipio(df, ano_completo, ultimo_mes_disponivel):
    """Agrupa valores por município."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]
    dados_por_municipio = df_filtered.groupby('CO_MUN')['VL_FOB'].sum().sort_values(ascending=False)
    total_municipios = dados_por_municipio.sum()
    return dados_por_municipio, total_municipios

def agregar_dados_por_produto(df, df_ncm, ano_completo, ultimo_mes_disponivel):
    """Agrupa produtos e calcula participação."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]

    def get_sh4(co_ncm):
        """Extrai SH4 considerando NCMs de 7 e 8 dígitos."""
        co_ncm_str = str(co_ncm)
        if pd.isna(co_ncm_str):
            return None
        if len(co_ncm_str) == 8:
            return co_ncm_str[:4]
        elif len(co_ncm_str) >= 7:
            return co_ncm_str[:3]
        else:
            return None

    df_filtered['SH4'] = df_filtered['CO_NCM'].apply(get_sh4).astype(str)
    df_sh4_not_null = df_filtered.dropna(subset=['SH4'])
    produtos = df_sh4_not_null.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(5)

    produtos_nomes = {}
    for sh4_code, valor in produtos.items():
        filtro_ncm = df_ncm[df_ncm['CO_SH4'] == sh4_code]
        if not filtro_ncm.empty:
            nome_produto = filtro_ncm['NO_SH4_POR'].iloc[0]
            produtos_nomes[nome_produto] = valor
        else:
            produtos_nomes[f"Produto NCM {sh4_code} não encontrado"] = valor

    return produtos_nomes


# --- FUNÇÕES DE IA (Refatoradas para aceitar api_key) ---

@st.cache_data
def obter_artigo_pais_gemini(nome_pais, api_key):
    """Chama a API do Gemini para obter o artigo de um país."""
    if not api_key:
        st.warning("Função de Artigo: API Key do Gemini não configurada nos 'Secrets'.")
        return None
        
    st.info(f"Consultando IA para obter o artigo de '{nome_pais}'...")
    
    # Prompt melhorado para artigos
    prompt = f"""Qual o artigo definido (o, a, os, as) correto para se referir ao país "{nome_pais}"? 
    Responda APENAS com o artigo.
    Por exemplo:
    - Para "Brasil" responda "o"
    - Para "China" responda "a"
    - Para "Estados Unidos" responda "os"
    - Para "Filipinas" responda "as"
    """

    url = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key=' + api_key
    headers = {'Content-Type': 'application/json'}
    data = {'contents': [{'parts': [{'text': prompt}]}]}

    try:
        resposta = requests.post(url, headers=headers, json=data, timeout=20)
        resposta.raise_for_status()

        conteudo_resposta = resposta.json()
        if 'candidates' in conteudo_resposta and conteudo_resposta['candidates']:
            texto_bruto = conteudo_resposta['candidates'][0]['content']['parts'][0]['text']
            artigo = texto_bruto.strip().replace('.', '').lower()
            if artigo in ['o', 'a', 'os', 'as']:
                return artigo
            else:
                st.warning(f"IA retornou um artigo inválido ('{artigo}') para '{nome_pais}'.")
                return None
        else:
            st.warning("A API não retornou um candidato válido para o artigo.")
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"Erro na chamada da API para obter artigo: {e}")
        return None


@st.cache_data
def chamar_gemini(texto, api_key):
    """Chama a API do Google Gemini para processar o texto."""
    if not api_key:
        st.warning("Função de Revisão: API Key do Gemini não configurada nos 'Secrets'.")
        return [texto] # Retorna o texto original em parágrafos
        
    url = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key=' + api_key
    headers = {'Content-Type': 'application/json'}
    data = {'contents': [{'parts': [{'text': texto}]}]}

    try:
        resposta = requests.post(url, headers=headers, json=data, timeout=60)
        resposta.raise_for_status()

        conteudo_resposta = resposta.json()
        if conteudo_resposta.get('candidates'):
            texto_processado = conteudo_resposta['candidates'][0]['content']['parts'][0]['text']
            paragraphs = texto_processado.split('\n\n')
            return [p.strip() for p in paragraphs if p.strip()]
        else:
            st.warning("A API não retornou nenhuma informação. Retornando texto original.")
            return [texto]
    except requests.exceptions.RequestException as e:
        st.error(f"Erro na chamada da API principal: {e}")
        return [texto]

# --- FUNÇÕES DE FORMATAÇÃO E UTILITÁRIOS ---

def formatar_valor(valor):
    prefixo = ""
    if valor < 0:
        prefixo = "-"
        valor = abs(valor)

    # Bilhões
    if valor >= 999_500_000:
        valor_em_bilhoes = round(valor / 1_000_000_000, 2)
        valor_formatado_str = f"{valor_em_bilhoes:,.2f}"
        unidade = "bilhão" if valor_em_bilhoes < 2 else "bilhões"
        resultado = f"US$ {valor_formatado_str} {unidade}"
    # Milhões
    elif valor >= 999_500:
        valor_em_milhoes = round(valor / 1_000_000, 2)
        valor_formatado_str = f"{valor_em_milhoes:,.2f}"
        unidade = "milhão" if valor_em_milhoes < 2 else "milhões"
        resultado = f"US$ {valor_formatado_str} {unidade}"
    # Mil
    elif valor >= 1_000:
        valor_formatado_str = f"{valor / 1_000:,.1f}"
        resultado = f"US$ {valor_formatado_str} mil"
    # Menor que 1.000
    else:
        resultado = f"US$ {valor:.0f}"

    resultado = resultado.replace(',', 'X').replace('.', ',').replace('X', '.')
    return f"{prefixo}{resultado}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

# --- CLASSE DE DOCUMENTO (Refatorada para aceitar caminhos) ---

class DocumentoApp:
    def __init__(self, logo_path):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""
        self.logo_path = logo_path
        # O diretório base agora é fixo para /tmp/ no servidor
        self.diretorio_base = "/tmp/" 

    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_conteudo_formatado(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_conteudo_central(self, texto):
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_paragrafo(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

D.adicionar_titulo(texto)
        p = self.doc.add_paragraph()
        if self.subsecao_atual == 0:
            run = p.add_run(f"{self.secao_atual}. {texto}")
        else:
            run = p.add_run(f"{self.secao_atual}.{self.subsecao_atual}. {texto}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def nova_secao(self):
        self.secao_atual += 1
        self.subsecao_atual = 0

    def nova_subsecao(self):
        self.subsecao_atual += 1

    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        largura_total_cm = 15.88
        table = header.add_table(rows=1, cols=2, width=Cm(largura_total_cm))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(2.91)
        table.columns[1].width = Cm(12.97)
        cell_imagem = table.cell(0, 0)
        paragraph_imagem = cell_imagem.paragraphs[0]
        paragraph_imagem.paragraph_format.space_before = Pt(0)
        paragraph_imagem.paragraph_format.space_after = Pt(0)
        
        run_imagem = paragraph_imagem.add_run()
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                run_imagem.add_picture(self.logo_path,
                                       width=Inches(1.71 / 2.54),
                                       height=Inches(1.67 / 2.54))
            except Exception as e:
                st.error(f"Erro ao adicionar imagem do logo ao Docx: {e}")
                paragraph_imagem.add_run("[Logo não encontrado]")
        else:
            paragraph_imagem.add_run("[Logo não encontrado]")

        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        
        # --- Formatação de cabeçalho explícita ---
        
        def formatar_paragrafo_cabecalho(p):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Linha 1: "GOVERNO DO ESTADO..."
        p = cell_texto.paragraphs[0]
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[0])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True # Linha 1: BOLD

        # Linha 2: "SECRETARIA DE ESTADO..."
        p = cell_texto.add_paragraph()
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[1])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True # Linha 2: BOLD
        
        # Linhas 3 e 4 (Subsecretaria, Superintendência)
        for texto in textos[2:]: 
            p = cell_texto.add_paragraph()
            formatar_paragrafo_cabecalho(p)
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = False # Linhas 3 e 4: REGULAR
        # --- Fim da formatação ---

    def finalizar_documento(self):
        """Salva o documento em memória e retorna."""
        
        diretorio_real = self.diretorio_base
        try:
            os.makedirs(diretorio_real, exist_ok=True)
        except Exception:
            # Em ambientes serverless, /tmp/ é o único local gravável
            diretorio_real = "/tmp/"
            os.makedirs(diretorio_real, exist_ok=True)
            
        nome_arquivo = f"{self.titulo_doc}.docx"
        nome_arquivo_sanitizado = sanitize_filename(nome_arquivo)
        caminho_completo = os.path.join(diretorio_real, nome_arquivo_sanitizado)

        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        
        file_bytes = file_stream.getvalue()
        st.success(f"Documento '{nome_arquivo_sanitizado}' gerado com sucesso!")
        
        try:
            self.doc.save(caminho_completo)
            st.info(f"Salvo no servidor em: {caminho_completo}")
        except Exception:
            pass 

        return file_bytes, nome_arquivo_sanitizado


# --- ----------------------------------- ---
# --- INTERFACE GRÁFICA DO STREAMLIT (NOVO "MAIN") ---
# --- ----------------------------------- ---

st.set_page_config(page_title="Gerador de Briefings ComexStat", layout="wide")
st.title(" automação de Briefings ComexStat")

# --- Inicialização do Session State ---
if 'arquivos_gerados' not in st.session_state:
    st.session_state.arquivos_gerados = []

# --- ENTRADAS DO USUÁRIO NA SIDEBAR ---

logo_sidebar_path = "LogoMinasGerais.png"
if os.path.exists(logo_sidebar_path):
    st.sidebar.image(logo_sidebar_path, width=150)

st.sidebar.header(" Configurações Avançadas")

# --- ALTERAÇÃO AQUI: Leitura da API Key (sem aviso) ---
api_key_ui = st.secrets.get("GEMINI_API_KEY") 
# --- FIM DA ALTERAÇÃO ---

# --- REMOVIDO: Input do diretório ---

revisao_texto_gemini_ui = st.sidebar.checkbox("Usar revisão de IA (Gemini)", value=False)


# --- ENTRADAS PRINCIPAIS ---
st.header("1. Configurações da Análise")

col1, col2 = st.columns(2)
with col1:
    ano_atual = datetime.now().year
    ano_selecionado = st.number_input(
        "Digite o ano:",
        min_value=1998,
        max_value=ano_atual,
        value=ano_atual - 1,
        help="Os dados de 1997 podem estar incompletos."
    )
with col2:
    paises_input = st.text_input(
        "País(es) (separados por '; '):",
        "China; Estados Unidos",
        help="Ex: 'China' ou 'China; Estados Unidos; Argentina'"
    )
    paises = [pais.strip() for pais in paises_input.split('; ')]

# --- LÓGICA CONDICIONAL PARA ENTRADAS ---
agrupado = False
nome_agrupamento = None

if len(paises) > 1:
    st.header("2. Opções de Agrupamento")
    agrupamento_input = st.radio(
        "Deseja que os dados sejam agrupados ou separados?",
        ("agrupados", "separados"),
        index=0,
        horizontal=True
    )
    agrupado = (agrupamento_input == "agrupados")

    if agrupado:
        quer_nome_agrupamento = st.checkbox("Deseja dar um nome para este agrupamento?")
        if quer_nome_agrupamento:
            nome_agrupamento = st.text_input("Digite o nome do agrupamento:")

# --- EXECUÇÃO DO SCRIPT ---
st.header("3. Gerar Relatório")

if st.button(" Iniciar Geração do Relatório"):
    
    st.session_state.arquivos_gerados = []
    
    logo_path_to_use = "LogoMinasGerais.png" 
    if not os.path.exists(logo_path_to_use):
        st.warning(f"Aviso: A logo 'LogoMinasGerais.png' não foi encontrada. O cabeçalho será gerado sem a logo.")
        logo_path_to_use = None
    
    with st.spinner(f"Gerando relatório para {paises_input} ({ano_selecionado})... Isso pode levar alguns minutos."):
        
        try:
            # --- Validação de Países ---
            codigos_paises, nomes_paises_validos, paises_invalidos = validar_paises(paises)
            if paises_invalidos:
                st.warning(f"Países não encontrados ou inválidos (ignorados): {', '.join(paises_invalidos)}")
            if not nomes_paises_validos:
                st.error("Nenhum país válido fornecido. A geração foi interrompida.")
                st.stop()
            
            # --- URLs ---
            url_exp_ano = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_selecionado}.csv"
            url_exp_ano_anterior = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_selecionado - 1}.csv"
            url_imp_ano = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_selecionado}.csv"
            url_imp_ano_anterior = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_selecionado - 1}.csv"
            url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
            url_exp_mun = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_selecionado}_MUN.csv"
            url_imp_mun = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_selecionado}_MUN.csv"
            url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
            
            # --- 1. Carregar dados comuns (pequenos) ---
            st.info("Carregando tabelas auxiliares (NCM, UF)...")
            df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv", usecols=['CO_SH4', 'NO_SH4_POR'])
            df_uf_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['CO_MUN_GEO', 'NO_MUN_MIN'])
            
            if df_ncm is None or df_uf_mun is None:
                st.error("Não foi possível carregar tabelas auxiliares. Abortando.")
                st.stop()

            # --- 2. Bloco de Exportação ---
            st.info("Processando dados de Exportação (NCM)...")
            df_exp_ano = carregar_dataframe(url_exp_ano, f"EXP_{ano_selecionado}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            df_exp_ano_anterior = carregar_dataframe(url_exp_ano_anterior, f"EXP_{ano_selecionado - 1}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)

            if df_exp_ano is None or df_exp_ano_anterior is None:
                st.error("Não foi possível carregar dados de exportação. Abortando.")
                st.stop()

            ultimo_mes_disponivel = df_exp_ano['CO_MES'].max()
            ano_completo = ultimo_mes_disponivel == 12

            df_exp_ano_estados = filtrar_dados_por_estado_e_mes(df_exp_ano, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
            df_exp_ano_anterior_estados = filtrar_dados_por_estado_e_mes(df_exp_ano_anterior, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
            df_exp_ano_mg = filtrar_dados_por_estado_e_mes(df_exp_ano, ['MG'], ultimo_mes_disponivel, ano_completo)
            df_exp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)
            df_exp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)
            exportacao_pais_ano = df_exp_ano_mg_paises['VL_FOB'].sum()
            exportacao_pais_ano_anterior = df_exp_ano_anterior_mg_paises['VL_FOB'].sum()
            
            if not agrupado:
                ranking_mg_exp = calcular_ranking_por_pais(df_exp_ano_mg, ano_completo, ultimo_mes_disponivel)
                try:
                    posicao_pais_mg_exp = ranking_mg_exp.index.get_loc(codigos_paises[0]) + 1
                except KeyError:
                    posicao_pais_mg_exp = 0
            else:
                posicao_pais_mg_exp = None
            
            df_exp_ano_mg_filtered = df_exp_ano_mg
            if not ano_completo:
                df_exp_ano_mg_filtered = df_exp_ano_mg[df_exp_ano_mg['CO_MES'] <= ultimo_mes_disponivel]
            exportacao_mg_total_ano = df_exp_ano_mg_filtered['VL_FOB'].sum()
            del df_exp_ano_mg_filtered 

            participacao_pais_mg_exp = calcular_participacao(exportacao_pais_ano, exportacao_mg_total_ano)
            diferenca_exportacao, tipo_diferenca_exp = calcular_diferenca_percentual(exportacao_pais_ano, exportacao_pais_ano_anterior, ano_completo, ultimo_mes_disponivel)
            
            df_exp_brasil_periodo = df_exp_ano
            if not ano_completo:
                df_exp_brasil_periodo = df_exp_ano[df_exp_ano['CO_MES'] <= ultimo_mes_disponivel]
            exportacao_mg_para_pais = df_exp_brasil_periodo[(df_exp_brasil_periodo['SG_UF_NCM'] == 'MG') & (df_exp_brasil_periodo['CO_PAIS'].isin(codigos_paises))]['VL_FOB'].sum()
            exportacao_brasil_pais = df_exp_brasil_periodo[df_exp_brasil_periodo['CO_PAIS'].isin(codigos_paises)]['VL_FOB'].sum()
            del df_exp_brasil_periodo 
            
            participacao_mg_brasil_exp = calcular_participacao(exportacao_mg_para_pais, exportacao_brasil_pais)
            posicao_mg_pais_exp = calcular_posicao_estado_pais(df_exp_ano_estados, codigos_paises, ano_completo, ultimo_mes_disponivel)
            produtos_exportacao = agregar_dados_por_produto(df_exp_ano_mg_paises.copy(), df_ncm, ano_completo, ultimo_mes_disponivel)
            
            st.info("Processando dados de Exportação (Municípios)...")
            df_exp_mun = carregar_dataframe(url_exp_mun, f"EXP_{ano_selecionado}_MUN.csv", usecols=MUN_COLS)
            if df_exp_mun is None:
                st.error("Não foi possível carregar dados de exportação por município. Abortando.")
                st.stop()
                
            df_exp_mun_filtrado = df_exp_mun[(df_exp_mun['SG_UF_MUN'] == 'MG') & (df_exp_mun['CO_PAIS'].isin(codigos_paises))]
            if not ano_completo:
                df_exp_mun_filtrado = df_exp_mun_filtrado[df_exp_mun_filtrado['CO_MES'] <= ultimo_mes_disponivel]
            exportacoes_por_municipio, total_exportacoes_municipios = agregar_dados_por_municipio(df_exp_mun_filtrado, ano_completo, ultimo_mes_disponivel)
            
            st.info("Liberando memória de exportação...")
            del df_exp_ano, df_exp_ano_anterior, df_exp_ano_estados, df_exp_ano_mg, df_exp_ano_mg_paises, df_exp_mun, df_exp_mun_filtrado
            
            # --- 4. Bloco de Importação ---
            st.info("Processando dados de Importação (NCM)...")
            df_imp_ano = carregar_dataframe(url_imp_ano, f"IMP_{ano_selecionado}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            df_imp_ano_anterior = carregar_dataframe(url_imp_ano_anterior, f"IMP_{ano_selecionado - 1}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            
            if df_imp_ano is None or df_imp_ano_anterior is None:
                st.error("Não foi possível carregar dados de importação. Abortando.")
                st.stop()
            
            df_imp_ano_estados = filtrar_dados_por_estado_e_mes(df_imp_ano, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
            df_imp_ano_anterior_estados = filtrar_dados_por_estado_e_mes(df_imp_ano_anterior, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
            df_imp_ano_mg = filtrar_dados_por_estado_e_mes(df_imp_ano, ['MG'], ultimo_mes_disponivel, ano_completo)
            df_imp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)
            df_imp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)
            importacao_pais_ano = df_imp_ano_mg_paises['VL_FOB'].sum()
            importacao_pais_ano_anterior = df_imp_ano_anterior_mg_paises['VL_FOB'].sum()
            
            if not agrupado:
                ranking_mg_imp = calcular_ranking_por_pais(df_imp_ano_mg, ano_completo, ultimo_mes_disponivel)
                try:
                    posicao_pais_mg_imp = ranking_mg_imp.index.get_loc(codigos_paises[0]) + 1
                except KeyError:
                    posicao_pais_mg_imp = 0
            else:
                posicao_pais_mg_imp = None
            
            df_imp_ano_mg_filtered = df_imp_ano_mg
            if not ano_completo:
                df_imp_ano_mg_filtered = df_imp_ano_mg[df_imp_ano_mg['CO_MES'] <= ultimo_mes_disponivel]
            importacao_mg_total_ano = df_imp_ano_mg_filtered['VL_FOB'].sum()
            del df_imp_ano_mg_filtered 

            participacao_pais_mg_imp = calcular_participacao(importacao_pais_ano, importacao_mg_total_ano)
            diferenca_importacao, tipo_diferenca_imp = calcular_diferenca_percentual(importacao_pais_ano, importacao_pais_ano_anterior, ano_completo, ultimo_mes_disponivel)
            
            df_imp_brasil_periodo = df_imp_ano
            if not ano_completo:
                df_imp_brasil_periodo = df_imp_ano[df_imp_ano['CO_MES'] <= ultimo_mes_disponivel]
            importacao_mg_para_pais = df_imp_brasil_periodo[(df_imp_brasil_periodo['SG_UF_NCM'] == 'MG') & (df_imp_brasil_periodo['CO_PAIS'].isin(codigos_paises))]['VL_FOB'].sum()
            importacao_brasil_pais = df_imp_brasil_periodo[df_imp_brasil_periodo['CO_PAIS'].isin(codigos_paises)]['VL_FOB'].sum()
            del df_imp_brasil_periodo 
            
            participacao_mg_brasil_imp = calcular_participacao(importacao_mg_para_pais, importacao_brasil_pais)
            posicao_mg_pais_imp = calcular_posicao_estado_pais(df_imp_ano_estados, codigos_paises, ano_completo, ultimo_mes_disponivel)
            produtos_importacao = agregar_dados_por_produto(df_imp_ano_mg_paises.copy(), df_ncm, ano_completo, ultimo_mes_disponivel)
            
            st.info("Processando dados de Importação (Municípios)...")
            df_imp_mun = carregar_dataframe(url_imp_mun, f"IMP_{ano_selecionado}_MUN.csv", usecols=MUN_COLS)
            if df_imp_mun is None:
                st.error("Não foi possível carregar dados de importação por município. Abortando.")
                st.stop()

            df_imp_mun_filtrado = df_imp_mun[(df_imp_mun['SG_UF_MUN'] == 'MG') & (df_imp_mun['CO_PAIS'].isin(codigos_paises))]
            if not ano_completo:
                df_imp_mun_filtrado = df_imp_mun_filtrado[df_imp_mun_filtrado['CO_MES'] <= ultimo_mes_disponivel]
            importacoes_por_municipio, total_importacoes_municipios = agregar_dados_por_municipio(df_imp_mun_filtrado, ano_completo, ultimo_mes_disponivel)
            
            st.info("Liberando memória de importação...")
            del df_imp_ano, df_imp_ano_anterior, df_imp_ano_estados, df_imp_ano_mg, df_imp_ano_mg_paises, df_imp_mun, df_imp_mun_filtrado

            # --- 6. Bloco de Cálculo Final (Balança/Fluxo) ---
            st.info("Calculando balança comercial...")
            balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo = calcular_balanca_e_fluxo(exportacao_pais_ano, importacao_pais_ano, exportacao_pais_ano_anterior, importacao_pais_ano_anterior)
            
            # --- 7. Geração de Texto e Documento ---
            st.info("Gerando documento .docx...")

            if agrupado:
                # --- LÓGICA PARA AGRUPADOS ---
                app = DocumentoApp(logo_path=logo_path_to_use)
                paises_corretos = nomes_paises_validos 
                nome_relatorio = nome_agrupamento if nome_agrupamento else ', '.join(paises_corretos)

                # --- Geração de Texto ... ---
                if ano_completo:
                    fluxo_e_balanca = f"Em {ano_selecionado}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, representando {'aumento' if variacao_fluxo > 0 else 'queda'} de {abs(variacao_fluxo):.2f}% em comparação a {ano_selecionado-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação a {ano_selecionado-1}."
                    frase_1 = fluxo_e_balanca
                else:
                    fluxo_e_balanca = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, com {'um aumento' if variacao_fluxo > 0 else 'uma queda'} de {abs(variacao_fluxo):.2f}% em comparação ao mesmo período em {ano_selecionado-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação ao mesmo período em {ano_selecionado-1}."
                    frase_1 = fluxo_e_balanca

                if ano_completo:
                    texto_exportacao = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano_selecionado}, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação a {ano_selecionado-1}. A participação {nome_relatorio} nas exportações totais de Minas Gerais em {ano_selecionado} foi equivalente a {participacao_pais_mg_exp}%. "
                    frase_2 = texto_exportacao
                    texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano_selecionado}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao(s) país(es)."
                    frase_3 = texto_exportacao_2
                else:
                    texto_exportacao = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} neste período, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação ao mesmo período em {ano_selecionado-1}. A participação {nome_relatorio} nas exportações totais de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, foi equivalente a {participacao_pais_mg_exp}%. "
                    frase_2 = texto_exportacao
                    texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao(s) país(es)."
                    frase_3 = texto_exportacao_2

                if ano_completo:
                    texto_produtos_exportacao = f"Em {ano_selecionado}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                    frase_6 = texto_produtos_exportacao
                else:
                    texto_produtos_exportacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                    frase_6 = texto_produtos_exportacao
                texto_produtos_exportacao_lista = []
                frase_6_lista = []
                for nome_produto, valor_fob in produtos_exportacao.items(): 
                    participacao_produto_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) 
                    texto_produtos_exportacao_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")
                    frase_6_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")
                texto_produtos_exportacao += "; ".join(texto_produtos_exportacao_lista) + "."
                frase_6 += "; ".join(frase_6_lista) + "."

                if ano_completo:
                    texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano_selecionado}, os principais foram: "
                    frase_7 = texto_municipios_exportacao
                else:
                    texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais foram: "
                    frase_7 = texto_municipios_exportacao
                texto_municipios_exportacao_lista = []
                frase_7_lista = []
                for i, (codigo_municipio, valor_fob) in enumerate(exportacoes_por_municipio.head(5).items()):
                    nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                    participacao_municipio_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano)
                    texto_municipios_exportacao_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")
                    frase_7_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")
                texto_municipios_exportacao += "; ".join(texto_municipios_exportacao_lista) + "."
                frase_7 += "; ".join(frase_7_lista) + "."

                if ano_completo:
                    texto_importacao = f"As importações mineiras provenientes {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} em {ano_selecionado}, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação a {ano_selecionado-1}. A participação {nome_relatorio} nas importações totais de Minas Gerais em {ano_selecionado} foi equivalente a {participacao_pais_mg_imp}%. "
                    frase_4 = texto_importacao
                    texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio} em {ano_selecionado}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao(s) país(es)."
                    frase_5 = texto_importacao_2
                else:
                    texto_importacao = f"As importações mineiras provenientes {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} neste período, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação ao mesmo período em {ano_selecionado-1}. A participação {nome_relatorio} nas importações totais de Minas Gerais em {ano_selecionado} foi equivalente a {participacao_pais_mg_imp}%. "
                    frase_4 = texto_importacao
                    texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao(s) país(es)."
                    frase_5 = texto_importacao_2

                if ano_completo:
                    texto_produtos_importacao = f"Em {ano_selecionado}, os principais produtos importados para Minas Gerais {nome_relatorio} foram: "
                    frase_8 = texto_produtos_importacao
                else:
                    texto_produtos_importacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais produtos importados para Minas Gerais {nome_relatorio} foram: "
                    frase_8 = texto_produtos_importacao
                texto_produtos_importacao_lista = []
                frase_8_lista = []
                for nome_produto, valor_fob in produtos_importacao.items():
                    participacao_produto_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                    texto_produtos_importacao_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")
                    frase_8_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")
                texto_produtos_importacao += "; ".join(texto_produtos_importacao_lista) + "."
                frase_8 += "; ".join(frase_8_lista) + "."

                if ano_completo:
                    texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio} em {ano_selecionado}, os principais foram: "
                    frase_9 = texto_municipios_importacao
                else:
                    texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais foram: "
                    frase_9 = texto_municipios_importacao
                texto_municipios_importacao_lista = []
                frase_9_lista = []
                for i, (codigo_municipio, valor_fob) in enumerate(importacoes_por_municipio.head(5).items()):
                    nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                    participacao_municipio_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                    texto_municipios_importacao_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")
                    frase_9_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")
                texto_municipios_importacao += "; ".join(texto_municipios_importacao_lista) + "."
                frase_9 += "; ".join(frase_9_lista) + "."

                # --- GEMINI ---
                texto_relatorio = frase_1 + '\n' + frase_2 + frase_3 + frase_6 + frase_7 + '\n' + frase_4 + frase_5 + frase_8 + frase_9
                texto_processado_ia_paragraphs = []
                
                if revisao_texto_gemini_ui:
                    st.info("Chamando IA para revisar o texto...")
                    prompt_gemini = f"Agrupe todos os pontos em 5 parágrafos, relacionando por assunto. Você não pode suprimir nenhuma das informações e não pode adicionar nenhuma palavra ou texto que forneça qualquer tipo de valoração ou juízo de valor. Ou seja, sua função é apenas transformar o texto de tópicos para parágrafos. A seguir, o texto: \n{texto_relatorio}"
                    texto_processado_ia_paragraphs = chamar_gemini(prompt_gemini, api_key_ui)
                
                # --- Montagem do Documento ---
                if nome_agrupamento:
                    titulo_documento = f"Briefing - {nome_agrupamento} - {ano_selecionado}"
                else:
                    titulo_documento = f"Briefing - {', '.join(paises_corretos)} - {ano_selecionado}"
                
                app.set_titulo(titulo_documento)
                app.nova_secao()
                app.adicionar_titulo("Fluxo Comercial")
                app.adicionar_conteudo_formatado(fluxo_e_balanca)
                app.nova_secao()
                app.adicionar_titulo("Exportações")
                app.adicionar_conteudo_formatado(texto_exportacao)
                app.adicionar_conteudo_formatado(texto_exportacao_2)
                app.adicionar_conteudo_formatado(texto_produtos_exportacao)
                app.adicionar_conteudo_formatado(texto_municipios_exportacao)
                app.nova_secao()
                app.adicionar_titulo("Importações")
                app.adicionar_conteudo_formatado(texto_importacao)
                app.adicionar_conteudo_formatado(texto_importacao_2)
                app.adicionar_conteudo_formatado(texto_produtos_importacao)
                app.adicionar_conteudo_formatado(texto_municipios_importacao)
                
                if texto_processado_ia_paragraphs: 
                    app.nova_secao()
                    app.adicionar_titulo("Texto processado pela IA")
                    for paragraph in texto_processado_ia_paragraphs:
                        app.adicionar_conteudo_formatado(paragraph)
                
                # Finaliza e prepara para download
                file_bytes, file_name = app.finalizar_documento() # Agora recebe bytes
                
                # Salva os bytes no state
                st.session_state.arquivos_gerados.append({"name": file_name, "data": file_bytes})

            else:
                # --- LÓGICA PARA SEPARADOS ---
                paises_corretos = nomes_paises_validos
                
                for pais in paises_corretos:
                    st.subheader(f"Processando: {pais}")
                    app = DocumentoApp(logo_path=logo_path_to_use)
                    
                    codigos_paises_loop = [obter_codigo_pais(pais)]

                    # --- 2. Bloco de Exportação (Separado) ---
                    st.info(f"Processando Exportação (NCM) para {pais}...")
                    df_exp_ano = carregar_dataframe(url_exp_ano, f"EXP_{ano_selecionado}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
                    df_exp_ano_anterior = carregar_dataframe(url_exp_ano_anterior, f"EXP_{ano_selecionado - 1}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
                    if df_exp_ano is None or df_exp_ano_anterior is None:
                        st.error(f"Não foi possível carregar dados de exportação para {pais}.")
                        continue

                    ultimo_mes_disponivel = df_exp_ano['CO_MES'].max()
                    ano_completo = ultimo_mes_disponivel == 12
                    
                    df_exp_ano_estados = filtrar_dados_por_estado_e_mes(df_exp_ano, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
                    df_exp_ano_mg = filtrar_dados_por_estado_e_mes(df_exp_ano, ['MG'], ultimo_mes_disponivel, ano_completo)
                    df_exp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano, codigos_paises_loop, False, ultimo_mes_disponivel, ano_completo)
                    df_exp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior, codigos_paises_loop, False, ultimo_mes_disponivel, ano_completo)
                    exportacao_pais_ano = df_exp_ano_mg_paises['VL_FOB'].sum()
                    exportacao_pais_ano_anterior = df_exp_ano_anterior_mg_paises['VL_FOB'].sum()
                    
                    ranking_mg_exp = calcular_ranking_por_pais(df_exp_ano_mg, ano_completo, ultimo_mes_disponivel)
                    try:
                        posicao_pais_mg_exp = ranking_mg_exp.index.get_loc(codigos_paises_loop[0]) + 1
                    except KeyError:
                        posicao_pais_mg_exp = 0
                    
                    df_exp_ano_mg_filtered = df_exp_ano_mg
                    if not ano_completo:
                        df_exp_ano_mg_filtered = df_exp_ano_mg[df_exp_ano_mg['CO_MES'] <= ultimo_mes_disponivel]
                    exportacao_mg_total_ano = df_exp_ano_mg_filtered['VL_FOB'].sum()
                    del df_exp_ano_mg_filtered

                    participacao_pais_mg_exp = calcular_participacao(exportacao_pais_ano, exportacao_mg_total_ano)
                    diferenca_exportacao, tipo_diferenca_exp = calcular_diferenca_percentual(exportacao_pais_ano, exportacao_pais_ano_anterior, ano_completo, ultimo_mes_disponivel)
                    
                    df_exp_brasil_periodo = df_exp_ano
                    if not ano_completo:
                        df_exp_brasil_periodo = df_exp_ano[df_exp_ano['CO_MES'] <= ultimo_mes_disponivel]
                    exportacao_mg_para_pais = df_exp_brasil_periodo[(df_exp_brasil_periodo['SG_UF_NCM'] == 'MG') & (df_exp_brasil_periodo['CO_PAIS'].isin(codigos_paises_loop))]['VL_FOB'].sum()
                    exportacao_brasil_pais = df_exp_brasil_periodo[df_exp_brasil_periodo['CO_PAIS'].isin(codigos_paises_loop)]['VL_FOB'].sum()
                    del df_exp_brasil_periodo
                    
                    participacao_mg_brasil_exp = calcular_participacao(exportacao_mg_para_pais, exportacao_brasil_pais)
                    posicao_mg_pais_exp = calcular_posicao_estado_pais(df_exp_ano_estados, codigos_paises_loop, ano_completo, ultimo_mes_disponivel)
                    produtos_exportacao = agregar_dados_por_produto(df_exp_ano_mg_paises.copy(), df_ncm, ano_completo, ultimo_mes_disponivel)
                    
                    st.info(f"Processando Exportação (Municípios) para {pais}...")
                    df_exp_mun = carregar_dataframe(url_exp_mun, f"EXP_{ano_selecionado}_MUN.csv", usecols=MUN_COLS)
                    if df_exp_mun is None:
                        st.error(f"Não foi possível carregar dados de exportação por município para {pais}.")
                        continue
                        
                    df_exp_mun_filtrado = df_exp_mun[(df_exp_mun['SG_UF_MUN'] == 'MG') & (df_exp_mun['CO_PAIS'].isin(codigos_paises_loop))]
                    if not ano_completo:
                        df_exp_mun_filtrado = df_exp_mun_filtrado[df_exp_mun_filtrado['CO_MES'] <= ultimo_mes_disponivel]
                    exportacoes_por_municipio, total_exportacoes_municipios = agregar_dados_por_municipio(df_exp_mun_filtrado, ano_completo, ultimo_mes_disponivel)
                    
                    st.info(f"Liberando memória de exportação de {pais}...")
                    del df_exp_ano, df_exp_ano_anterior, df_exp_ano_estados, df_exp_ano_mg, df_exp_ano_mg_paises, df_exp_mun, df_exp_mun_filtrado

                    # --- 4. Bloco de Importação (Separado) ---
                    st.info(f"Processando Importação (NCM) para {pais}...")
                    df_imp_ano = carregar_dataframe(url_imp_ano, f"IMP_{ano_selecionado}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
                    df_imp_ano_anterior = carregar_dataframe(url_imp_ano_anterior, f"IMP_{ano_selecionado - 1}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
                    if df_imp_ano is None or df_imp_ano_anterior is None:
                        st.error(f"Não foi possível carregar dados de importação para {pais}.")
                        continue
                    
                    df_imp_ano_estados = filtrar_dados_por_estado_e_mes(df_imp_ano, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
                    df_imp_ano_mg = filtrar_dados_por_estado_e_mes(df_imp_ano, ['MG'], ultimo_mes_disponivel, ano_completo)
                    df_imp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano, codigos_paises_loop, False, ultimo_mes_disponivel, ano_completo)
                    df_imp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior, codigos_paises_loop, False, ultimo_mes_disponivel, ano_completo)
                    importacao_pais_ano = df_imp_ano_mg_paises['VL_FOB'].sum()
                    importacao_pais_ano_anterior = df_imp_ano_anterior_mg_paises['VL_FOB'].sum()
                    
                    ranking_mg_imp = calcular_ranking_por_pais(df_imp_ano_mg, ano_completo, ultimo_mes_disponivel)
                    try:
                        posicao_pais_mg_imp = ranking_mg_imp.index.get_loc(codigos_paises_loop[0]) + 1
                    except KeyError:
                        posicao_pais_mg_imp = 0
                    
                    df_imp_ano_mg_filtered = df_imp_ano_mg
                    if not ano_completo:
                        df_imp_ano_mg_filtered = df_imp_ano_mg[df_imp_ano_mg['CO_MES'] <= ultimo_mes_disponivel]
                    importacao_mg_total_ano = df_imp_ano_mg_filtered['VL_FOB'].sum()
                    del df_imp_ano_mg_filtered

                    participacao_pais_mg_imp = calcular_participacao(importacao_pais_ano, importacao_mg_total_ano)
                    diferenca_importacao, tipo_diferenca_imp = calcular_diferenca_percentual(importacao_pais_ano, importacao_pais_ano_anterior, ano_completo, ultimo_mes_disponivel)
                    
                    df_imp_brasil_periodo = df_imp_ano
                    if not ano_completo:
                        df_imp_brasil_periodo = df_imp_ano[df_imp_ano['CO_MES'] <= ultimo_mes_disponivel]
                    importacao_mg_para_pais = df_imp_brasil_periodo[(df_imp_brasil_periodo['SG_UF_NCM'] == 'MG') & (df_imp_brasil_periodo['CO_PAIS'].isin(codigos_paises_loop))]['VL_FOB'].sum()
                    importacao_brasil_pais = df_imp_brasil_periodo[df_imp_brasil_periodo['CO_PAIS'].isin(codigos_paises_loop)]['VL_FOB'].sum()
                    del df_imp_brasil_periodo
                    
                    participacao_mg_brasil_imp = calcular_participacao(importacao_mg_para_pais, importacao_brasil_pais)
                    posicao_mg_pais_imp = calcular_posicao_estado_pais(df_imp_ano_estados, codigos_paises_loop, ano_completo, ultimo_mes_disponivel)
                    produtos_importacao = agregar_dados_por_produto(df_imp_ano_mg_paises.copy(), df_ncm, ano_completo, ultimo_mes_disponivel)
                    
                    st.info(f"Processando Importação (Municípios) para {pais}...")
                    df_imp_mun = carregar_dataframe(url_imp_mun, f"IMP_{ano_selecionado}_MUN.csv", usecols=MUN_COLS)
                    if df_imp_mun is None:
                        st.error(f"Não foi possível carregar dados de importação por município para {pais}.")
                        continue
                        
                    df_imp_mun_filtrado = df_imp_mun[(df_imp_mun['SG_UF_MUN'] == 'MG') & (df_imp_mun['CO_PAIS'].isin(codigos_paises_loop))]
                    if not ano_completo:
                        df_imp_mun_filtrado = df_imp_mun_filtrado[df_imp_mun_filtrado['CO_MES'] <= ultimo_mes_disponivel]
                    importacoes_por_municipio, total_importacoes_municipios = agregar_dados_por_municipio(df_imp_mun_filtrado, ano_completo, ultimo_mes_disponivel)
                    
                    st.info(f"Liberando memória de importação de {pais}...")
                    del df_imp_ano, df_imp_ano_anterior, df_imp_ano_estados, df_imp_ano_mg, df_imp_ano_mg_paises, df_imp_mun, df_imp_mun_filtrado
                    
                    # --- 6. Cálculo Final (Separado) ---
                    st.info(f"Calculando balança para {pais}...")
                    balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo = calcular_balanca_e_fluxo(exportacao_pais_ano, importacao_pais_ano, exportacao_pais_ano_anterior, importacao_pais_ano_anterior)

                    # --- 7. Geração de Texto e Documento (Separado) ---
                    st.info(f"Gerando documento .docx para {pais}...")
                    
                    nome_pais_base = pais

                    # --- ARTIGO ---
                    nome_relatorio = nome_pais_base
                    nome_relatorio_capitalizado = nome_pais_base
                    artigo = obter_artigo_pais_gemini(nome_pais_base, api_key_ui)
                    valid_articles = ['o', 'a', 'os', 'as']

                    if artigo and artigo.lower() in valid_articles:
                        nome_relatorio = f"{artigo.lower()} {nome_pais_base}"
                        nome_relatorio_capitalizado = f"{artigo.capitalize()} {nome_pais_base}"
                    else:
                        artigo = None
                        # Aviso silencioso no log do servidor, não para o usuário
                        print(f"Não foi possível obter um artigo válido para '{nome_pais_base}'. Usando nome do país sem artigo.")


                    contracoes_map = {'o': 'do', 'a': 'da', 'os': 'dos', 'as': 'das'}
                    preposicao_contraida = contracoes_map.get(artigo)

                    if preposicao_contraida:
                        nome_relatorio_com_contracao = f"{preposicao_contraida} {nome_pais_base}"
                    else:
                        nome_relatorio_com_contracao = f"de {nome_pais_base}"

                    titulo_documento = f"Briefing - {nome_pais_base} - {ano_selecionado}"
                    
                    # --- Geração de Texto ... ---
                    if ano_completo:
                        fluxo_e_balanca = f"Em {ano_selecionado}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, representando {'aumento' if variacao_fluxo > 0 else 'queda'} de {abs(variacao_fluxo):.2f}% em comparação a {ano_selecionado-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação a {ano_selecionado-1}."
                        frase_1 = fluxo_e_balanca
                    else:
                        fluxo_e_balanca = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, com {'um aumento' if variacao_fluxo > 0 else 'uma queda'} de {abs(variacao_fluxo):.2f}% em comparação ao mesmo período em {ano_selecionado-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação ao mesmo período em {ano_selecionado-1}."
                        frase_1 = fluxo_e_balanca

                    if posicao_pais_mg_exp > 0: 
                        if ano_completo:
                            texto_exportacao = f"{nome_relatorio_capitalizado} foi o {posicao_pais_mg_exp}º destino das exportações de Minas Gerais em {ano_selecionado}. As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano_selecionado}, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação a {ano_selecionado-1}. A participação {nome_relatorio_com_contracao} nas exportações totais de Minas Gerais em {ano_selecionado} foi equivalente a {participacao_pais_mg_exp}%. "
                            frase_2 = texto_exportacao
                        else:
                            texto_exportacao = f"{nome_relatorio_capitalizado} foi o {posicao_pais_mg_exp}º destino das exportações de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}. As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} neste período, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação ao mesmo período em {ano_selecionado-1}. A participação {nome_relatorio_com_contracao} nas exportações totais de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, foi equivalente a {participacao_pais_mg_exp}%. "
                            frase_2 = texto_exportacao
                    else: 
                        if ano_completo:
                            texto_exportacao = f"Em {ano_selecionado}, Minas Gerais não registrou exportações para {nome_relatorio}."
                            frase_2 = texto_exportacao
                        else:
                            texto_exportacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, Minas Gerais não registrou exportações para {nome_relatorio}."
                            frase_2 = texto_exportacao

                    if posicao_mg_pais_exp > 0: 
                        if ano_completo:
                            texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano_selecionado}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao país."
                            frase_3 = texto_exportacao_2
                        else:
                            texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao país."
                            frase_3 = texto_exportacao_2
                    else: 
                        texto_exportacao_2 = f"O estado de Minas Gerais não se posicionou no ranking de exportadores brasileiros para {nome_relatorio}, pois não houve registro de vendas."
                        frase_3 = texto_exportacao_2

                    if ano_completo:
                        texto_produtos_exportacao = f"Em {ano_selecionado}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                        frase_6 = texto_produtos_exportacao
                    else:
                        texto_produtos_exportacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                        frase_6 = texto_produtos_exportacao
                    texto_produtos_exportacao_lista = []
                    frase_6_lista = []
                    for nome_produto, valor_fob in produtos_exportacao.items(): 
                        participacao_produto_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) 
                        texto_produtos_exportacao_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")
                        frase_6_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")
                    texto_produtos_exportacao += "; ".join(texto_produtos_exportacao_lista) + "."
                    frase_6 += "; ".join(frase_6_lista) + "."

                    if ano_completo:
                        texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano_selecionado}, os principais foram: "
                        frase_7 = texto_municipios_exportacao
                    else:
                        texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais foram: "
                        frase_7 = texto_municipios_exportacao
                    texto_municipios_exportacao_lista = []
                    frase_7_lista = []
                    for i, (codigo_municipio, valor_fob) in enumerate(exportacoes_por_municipio.head(5).items()):
                        nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                        participacao_municipio_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) 
                        texto_municipios_exportacao_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")
                        frase_7_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")
                    texto_municipios_exportacao += "; ".join(texto_municipios_exportacao_lista) + "."
                    frase_7 += "; ".join(frase_7_lista) + "."

                    if posicao_pais_mg_imp > 0: 
                        if ano_completo:
                            texto_importacao = f"{nome_relatorio_capitalizado} foi a {posicao_pais_mg_imp}ª origem das importações de Minas Gerais em {ano_selecionado}. As importações mineiras provenientes {nome_relatorio_com_contracao} somaram {formatar_valor(importacao_pais_ano)} em {ano_selecionado}, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação a {ano_selecionado-1}. A participação {nome_relatorio_com_contracao} nas importações totais de Minas Gerais em {ano_selecionado} foi equivalente a {participacao_pais_mg_imp}%. "
                            frase_4 = texto_importacao
                        else:
                            texto_importacao = f"{nome_relatorio_capitalizado} foi a {posicao_pais_mg_imp}ª origem das importações de Minas Gerais até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}. As importações mineiras provenientes {nome_relatorio_com_contracao} somaram {formatar_valor(importacao_pais_ano)} neste período, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação ao mesmo período em {ano_selecionado-1}. A participação {nome_relatorio_com_contracao} nas importações totais de Minas Gerais em {ano_selecionado} foi equivalente a {participacao_pais_mg_imp}%. "
                            frase_4 = texto_importacao
                    else: 
                        if ano_completo:
                            texto_importacao = f"Em {ano_selecionado}, Minas Gerais não registrou importações provenientes {nome_relatorio_com_contracao}."
                            frase_4 = texto_importacao
                        else:
                            texto_importacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, Minas Gerais não registrou importações provenientes {nome_relatorio_com_contracao}."
                            frase_4 = texto_importacao

                    if posicao_mg_pais_imp > 0: 
                        if ano_completo:
                            texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio_com_contracao} em {ano_selecionado}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao país."
                            frase_5 = texto_importacao_2
                        else:
                            texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio_com_contracao} no período até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao país."
                            frase_5 = texto_importacao_2
                    else: 
                        texto_importacao_2 = f"O estado de Minas Gerais não se posicionou no ranking de importadores brasileiros {nome_relatorio_com_contracao}, pois não houve registro de compras."
                        frase_5 = texto_importacao_2

                    if ano_completo:
                        texto_produtos_importacao = f"Em {ano_selecionado}, os principais produtos importados para Minas Gerais {nome_relatorio_com_contracao} foram: "
                        frase_8 = texto_produtos_importacao
                    else:
                        texto_produtos_importacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais produtos importados para Minas Gerais {nome_relatorio_com_contracao} foram: "
                        frase_8 = texto_produtos_importacao
                    texto_produtos_importacao_lista = []
                    frase_8_lista = []
                    for nome_produto, valor_fob in produtos_importacao.items():
                        participacao_produto_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                        texto_produtos_importacao_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")
                        frase_8_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")
                    texto_produtos_importacao += "; ".join(texto_produtos_importacao_lista) + "."
                    frase_8 += "; ".join(frase_8_lista) + "."


                    if ano_completo:
                        texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio_com_contracao} em {ano_selecionado}, os principais foram: "
                        frase_9 = texto_municipios_importacao
                    else:
                        texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio_com_contracao} até {meses_pt[ultimo_mes_disponivel]} de {ano_selecionado}, os principais foram: "
                        frase_9 = texto_municipios_importacao
                    texto_municipios_importacao_lista = []
                    frase_9_lista = []
                    for i, (codigo_municipio, valor_fob) in enumerate(importacoes_por_municipio.head(5).items()):
                        nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                        participacao_municipio_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                        texto_municipios_importacao_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")
                        frase_9_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")
                    texto_municipios_importacao += "; ".join(texto_municipios_importacao_lista) + "."
                    frase_9 += "; ".join(frase_9_lista) + "."

                    # --- GEMINI ---
                    texto_processado_ia_paragraphs = None
                    if revisao_texto_gemini_ui:
                      st.info(f"Chamando IA para revisar texto de {pais}...")
                      texto_relatorio = frase_1 + '\n' + frase_2 + '\n' + frase_3 + '\n' + frase_6 + '\n' + frase_7 + '\n' + frase_4 + '\n' + frase_5 + '\n' + frase_8 + '\n' + frase_9
                      prompt_gemini = f"Ajuste a ortografia e concordância das orações a seguir. Você não pode suprimir nenhuma das informações e não pode adicionar nenhuma palavra ou texto que forneça qualquer tipo de valoração ou juízo de valor. Ou seja, sua função é apenas fazer ajustes de ortografia e concordância nas orações, mantendo todas as informações. Faça o retorno em formatação simples. A seguir, as orações: \n{texto_relatorio}"
                      texto_processado_ia_paragraphs = chamar_gemini(prompt_gemini, api_key_ui)
                    
                    # --- Montagem do Documento ---
                    app.set_titulo(titulo_documento)
                    app.nova_secao()
                    app.adicionar_titulo("Fluxo Comercial")
                    app.adicionar_conteudo_formatado(fluxo_e_balanca)
                    app.nova_secao()
                    app.adicionar_titulo("Exportações")
                    app.adicionar_conteudo_formatado(texto_exportacao)
                    app.adicionar_conteudo_formatado(texto_exportacao_2)
                    app.adicionar_conteudo_formatado(texto_produtos_exportacao)
                    app.adicionar_conteudo_formatado(texto_municipios_importacao)
                    app.nova_secao()
                    app.adicionar_titulo("Importações")
                    app.adicionar_conteudo_formatado(texto_importacao)
                    app.adicionar_conteudo_formatado(texto_importacao_2)
                    app.adicionar_conteudo_formatado(texto_produtos_importacao)
                    app.adicionar_conteudo_formatado(texto_municipios_importacao)

                    if texto_processado_ia_paragraphs:
                        app.nova_secao()
                        app.adicionar_titulo("Texto processado pela IA")
                        for paragraph in texto_processado_ia_paragraphs:
                            app.adicionar_conteudo_formatado(paragraph)
                    
                    # Finaliza e salva na lista
                    file_bytes, file_name = app.finalizar_documento() # Agora recebe bytes
                    
                    # Salva os bytes no state
                    st.session_state.arquivos_gerados.append({"name": file_name, "data": file_bytes})
                
        except Exception as e:
            st.error(f"Ocorreu um erro inesperado durante a geração:")
            st.exception(e)

# --- ----------------------------------- ---
# --- Bloco de exibição de Download (COM LÓGICA DE ZIP) ---
# --- ----------------------------------- ---

if st.session_state.arquivos_gerados:
    st.header("4. Relatórios Gerados")
    st.info("Clique para baixar os relatórios. Eles permanecerão aqui até que você gere um novo relatório.")
    
    if len(st.session_state.arquivos_gerados) > 1:
        # Caso "Separados": Criar um ZIP
        st.subheader("Pacote de Relatórios (ZIP)")
        
        # Cria o zip em memória
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            # Adiciona cada arquivo do session_state ao zip
            for arquivo in st.session_state.arquivos_gerados:
                zip_file.writestr(arquivo["name"], arquivo["data"])
        
        zip_bytes = zip_buffer.getvalue()
        
        st.download_button(
            label=f"Baixar todos os {len(st.session_state.arquivos_gerados)} relatórios (.zip)",
            data=zip_bytes,
            file_name=f"Briefings_ComexStat_{ano_selecionado}.zip", # Usa o ano selecionado
            mime="application/zip",
            key="download_zip"
        )
        
    elif len(st.session_state.arquivos_gerados) == 1:
        # Caso "Agrupado": Botão único
        st.subheader("Relatório Gerado")
        arquivo = st.session_state.arquivos_gerados[0] # Pega o único arquivo
        st.download_button(
            label=f"Baixar Relatório ({arquivo['name']})",
            data=arquivo["data"], 
            file_name=arquivo["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{arquivo['name']}"
        )

# --- Adiciona o rodapé ---
st.footer("Desenvolvido por Aest - Dados e Subsecretaria de Promoção de Investimentos e Cadeias Produtivas")
