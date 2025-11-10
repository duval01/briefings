import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime, timezone
import calendar
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import json
from google.api_core.client_options import ClientOptions
from googleapiclient import discovery

requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)

cache_dfs = {}

api_key = 'AIzaSyCNONTudhs5cgcsOaRe-CQuwSdSmd_Bs7Y'

revisao_texto_gemini = False

diretorio = "/content/drive/MyDrive/Briefings - Comex Stat/Países/"

estados_brasileiros = {'AC', 'AL', 'AP', 'AM', 'BA', 'CE', 'DF', 'ES', 'GO', 'MA', 'MT', 'MS', 'MG', 'PA', 'PB', 'PR',
                      'PE', 'PI', 'RJ', 'RN', 'RS', 'RO', 'RR', 'SC', 'SE', 'SP', 'TO'}
meses_pt = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro"
}

def ler_dados_csv_online(url):
    retries = 3
    for attempt in range(retries):
        try:
            resposta = requests.get(url, verify=False, timeout=(10, 60))
            resposta.raise_for_status()
            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1',
                             sep=';', dtype={'CO_SH4': str, 'CO_NCM': str})
            return df
        except requests.exceptions.RequestException as e:
            print(f"Erro ao acessar o arquivo CSV (tentativa {attempt + 1}/{retries}): {e}")
            if attempt < retries - 1 and "IncompleteRead" in str(e):
                print("Retentando download...")
                continue
            else:
                print(f"Falha ao baixar após {retries} tentativas.")
                return None
        except Exception as e:
            print(f"Erro inesperado ao baixar ou processar o arquivo CSV: {e}")
            return None
    return None


def carregar_dataframe(url, nome_arquivo):
    """Carrega o DataFrame do cache, do disco ou da URL."""
    if url in cache_dfs:
        print(f"Carregando a base {nome_arquivo} do cache...")
        return cache_dfs[url]
    elif os.path.exists(nome_arquivo):
        data_modificacao_local = datetime.fromtimestamp(os.path.getmtime(nome_arquivo), tz=timezone.utc)
        resposta = requests.head(url, verify=False, timeout=10)
        resposta.raise_for_status()
        last_modified = resposta.headers.get('Last-Modified')
        if last_modified:
            data_modificacao_remota = datetime.strptime(last_modified,
                                                        '%a, %d %b %Y %H:%M:%S %Z').replace(tzinfo=timezone.utc)
        else:
            print(f"Não foi possível obter a data de modificação remota para {nome_arquivo}. Carregando versão local.")
            df = pd.read_csv(nome_arquivo, encoding='latin-1', sep=';', dtype={'CO_SH4': str, 'CO_NCM': str})
            cache_dfs[url] = df
            return df

        if data_modificacao_remota > data_modificacao_local:
            print(f"Arquivo remoto {nome_arquivo} atualizado. Baixando...")
            df = ler_dados_csv_online(url)
            if df is not None:
                df.to_csv(nome_arquivo, encoding='latin-1', sep=';', index=False)
                cache_dfs[url] = df
            return df
        else:
            print(f"Carregando a base {nome_arquivo} do disco...")
            df = pd.read_csv(nome_arquivo, encoding='latin-1', sep=';', dtype={'CO_SH4': str, 'CO_NCM': str})
            cache_dfs[url] = df
            return df
    else:
        print(f"Baixando {nome_arquivo} da URL...")
        df = ler_dados_csv_online(url)
        if df is not None:
            df.to_csv(nome_arquivo, encoding='latin-1', sep=';', index=False)
            cache_dfs[url] = df
        return df

def obter_codigo_pais(nome_pais):
    """Obtém o código do país a partir do nome."""
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv")
    if df_pais is not None and not df_pais.empty:  
        filtro_pais = df_pais[df_pais['NO_PAIS'] == nome_pais]
        if not filtro_pais.empty:  
            codigo_pais = filtro_pais['CO_PAIS'].iloc[0]
            return codigo_pais
    return None

def validar_paises(paises):
    """Valida e corrige a lista de países, e solicita nova entrada, se necessário."""
    codigos_paises = []
    nomes_paises_validos = []

    for pais in paises:
        codigo_pais = obter_codigo_pais(pais)
        while codigo_pais is None or pais.lower() == "brasil":
            if pais.lower() == "brasil":
                print("Não é possível fazer a busca no Brasil. Digite outro país:")
            else:
                print(f"O país '{pais}' não foi encontrado no banco de dados do ComexStat.")

            # Permite a entrada de múltiplos países novamente
            paises_input = input("Digite o(s) país(es) (separados por ponto e vírgula e espaço): ")
            paises_temp = [p.strip() for p in paises_input.split('; ')]

            # Atualiza a lista de países e o loop
            paises = paises_temp
            pais = paises[0]  # Define o país atual como o primeiro da nova lista
            codigo_pais = obter_codigo_pais(pais)

        codigos_paises.append(codigo_pais)
        nomes_paises_validos.append(pais)

    return codigos_paises, nomes_paises_validos

def filtrar_dados_por_estado_e_mes(df, estados, ultimo_mes_disponivel, ano_completo):
    """Filtra o DataFrame por estado e, se necessário, por mês."""
    df_filtrado = df[df['SG_UF_NCM'].isin(list(estados))]
    if not ano_completo:
        df_filtrado = df_filtrado[df_filtrado['CO_MES'] <= ultimo_mes_disponivel]
    return df_filtrado

def filtrar_dados_por_mg_e_pais(df, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo):
    """Filtra o DataFrame por MG e pelos países (se agrupado) ou por país individual (se não agrupado) e por mês se ano incompleto."""
    df_filtrado = df[df['SG_UF_NCM'] == 'MG']
    if agrupado:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'].isin(codigos_paises)]
    else:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'] == codigos_paises[0]]
    if not ano_completo:
        df_filtrado = df_filtrado[df_filtrado['CO_MES'] <= ultimo_mes_disponivel]
    return df_filtrado

def filtrar_dados_municipios(df_exp_mun, df_imp_mun, codigos_paises, ultimo_mes_disponivel, ano_completo):
    """Filtra os DataFrames de municípios por MG, país(es) e, se necessário, por mês."""
    df_exp_mun_filtrado = df_exp_mun[(df_exp_mun['SG_UF_MUN'] == 'MG') & (df_exp_mun['CO_PAIS'].isin(codigos_paises))]
    df_imp_mun_filtrado = df_imp_mun[(df_imp_mun['SG_UF_MUN'] == 'MG') & (df_imp_mun['CO_PAIS'].isin(codigos_paises))]
    if not ano_completo:
        df_exp_mun_filtrado = df_exp_mun_filtrado[df_exp_mun_filtrado['CO_MES'] <= ultimo_mes_disponivel]
        df_imp_mun_filtrado = df_imp_mun_filtrado[df_imp_mun_filtrado['CO_MES'] <= ultimo_mes_disponivel]
    return df_exp_mun_filtrado, df_imp_mun_filtrado

def calcular_soma_por_estado(df, ano, ano_completo, ultimo_mes_disponivel, df_anterior=None):
    """Calcula a soma dos valores por estado para o ano e, opcionalmente, para o ano anterior."""
    soma_ano = df.groupby('SG_UF_NCM')['VL_FOB'].sum()
    if df_anterior is not None:
        df_anterior_filtered = df_anterior
        if not ano_completo:
            df_anterior_filtered = df_anterior[df_anterior['CO_MES'] <= ultimo_mes_disponivel]
        soma_ano_anterior = df_anterior_filtered.groupby('SG_UF_NCM')['VL_FOB'].sum()
        return soma_ano, soma_ano_anterior
    else:
        return soma_ano

def calcular_classificacao_estados(soma_ano, soma_ano_anterior, ano):
    """Calcula a classificação dos estados por valor total."""
    classificacao = pd.concat([soma_ano, soma_ano_anterior], axis=1, keys=[str(ano), str(ano - 1)]).sort_values(
        by=str(ano), ascending=False)
    return classificacao

def calcular_posicao_mg(classificacao):
    """Calcula a posição de MG na classificação."""
    posicao_mg = classificacao.index.get_loc('MG') + 1
    return posicao_mg

def calcular_ranking_por_pais(df, ano_completo, ultimo_mes_disponivel):
    """Calcula a posição do país no ranking de MG."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]
    ranking = df_filtered.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
    return ranking

def calcular_participacao(valor_parcial, valor_total):
    """Calcula a participação percentual."""
    if valor_total == 0:
        return 0.0
    participacao = round(valor_parcial / valor_total * 100, 2)
    return participacao

def calcular_diferenca_percentual(valor_atual, valor_anterior, ano_completo, ultimo_mes_disponivel):
    """Calcula a diferença percentual entre dois valores."""
    if valor_anterior == 0:
        return 0.0, "acréscimo" if valor_atual > 0 else "redução" if valor_atual < 0 else "estabilidade"
    if ano_completo:
        diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    else:
        diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    if diferenca > 0:
        tipo_diferenca = "um acréscimo"
    elif diferenca < 0:
        tipo_diferenca = "uma redução"
    else:
        tipo_diferenca = "uma estabilidade"
    diferenca = abs(diferenca)
    return diferenca, tipo_diferenca

def calcular_posicao_estado_pais(df, codigos_paises, ano_completo, ultimo_mes_disponivel):
    """Calcula a posição de MG nas exportações/importações para o(s) país(es) informado(s)."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]

    df_comercio_pais = df_filtered[df_filtered['CO_PAIS'].isin(codigos_paises)]

    if df_comercio_pais.empty:
        return 0 

    ranking_estados_pais = df_comercio_pais.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(
        ascending=False)

    if 'MG' not in ranking_estados_pais.index:
        return 0 

    posicao_mg_pais = ranking_estados_pais.index.get_loc('MG') + 1
    return posicao_mg_pais

def calcular_balanca_e_fluxo(exportacao_ano, importacao_ano, exportacao_ano_anterior, importacao_ano_anterior):
    """Calcula a balança comercial, o fluxo comercial e suas variações."""
    balanca_ano = exportacao_ano - importacao_ano
    balanca_ano_anterior = exportacao_ano_anterior - importacao_ano_anterior
    fluxo_comercial_ano = exportacao_ano + importacao_ano
    fluxo_comercial_ano_anterior = exportacao_ano_anterior + importacao_ano_anterior
    variacao_balanca = 0
    variacao_fluxo = 0
    if balanca_ano_anterior != 0:
        variacao_balanca = ((balanca_ano - balanca_ano_anterior) / balanca_ano_anterior) * 100
    if fluxo_comercial_ano_anterior != 0:
        variacao_fluxo = ((fluxo_comercial_ano - fluxo_comercial_ano_anterior) / fluxo_comercial_ano_anterior) * 100

    return balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo

def agregar_dados_por_municipio(df, ano_completo, ultimo_mes_disponivel):
    """Agrupa valores por município."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]
    dados_por_municipio = df_filtered.groupby('CO_MUN')['VL_FOB'].sum().sort_values(ascending=False)
    total_municipios = dados_por_municipio.sum()
    return dados_por_municipio, total_municipios

def agregar_dados_por_produto(df, df_ncm, ano_completo, ultimo_mes_disponivel):
    """Agrupa produtos e calcula participação."""
    df_filtered = df
    if not ano_completo:
        df_filtered = df[df['CO_MES'] <= ultimo_mes_disponivel]

    def get_sh4(co_ncm):
        """Extrai SH4 considerando NCMs de 7 e 8 dígitos."""
        co_ncm_str = str(co_ncm)
        if pd.isna(co_ncm_str):
            return None
        if len(co_ncm_str) == 8:
            return co_ncm_str[:4]
        elif len(co_ncm_str) >= 7:
            return co_ncm_str[:3]
        else:
            return None

    df_filtered['SH4'] = df_filtered['CO_NCM'].apply(get_sh4).astype(str)
    df_sh4_not_null = df_filtered.dropna(subset=['SH4'])
    produtos = df_sh4_not_null.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(5)

    produtos_nomes = {}
    for sh4_code, valor in produtos.items():
        filtro_ncm = df_ncm[df_ncm['CO_SH4'] == sh4_code]
        if not filtro_ncm.empty:
            nome_produto = filtro_ncm['NO_SH4_POR'].iloc[0]
            produtos_nomes[nome_produto] = valor
        else:
            produtos_nomes[f"Produto NCM {sh4_code} não encontrado"] = valor

    return produtos_nomes

def calcular_fluxo_comercial(ano, paises, agrupado):
    """Calcula o fluxo comercial de Minas Gerais com um ou mais países em um determinado ano."""

    codigos_paises, nomes_paises_validos = validar_paises(paises)

    # URLs dos arquivos CSV do ComexStat
    url_exp_ano = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano}.csv"
    url_exp_ano_anterior = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano - 1}.csv"
    url_imp_ano = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano}.csv"
    url_imp_ano_anterior = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano - 1}.csv"
    url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
    url_exp_mun = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano}_MUN.csv"
    url_imp_mun = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano}_MUN.csv"
    url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"

    df_exp_ano = carregar_dataframe(url_exp_ano, f"EXP_{ano}.csv")
    df_exp_ano_anterior = carregar_dataframe(url_exp_ano_anterior, f"EXP_{ano - 1}.csv")
    df_imp_ano = carregar_dataframe(url_imp_ano, f"IMP_{ano}.csv")
    df_imp_ano_anterior = carregar_dataframe(url_imp_ano_anterior, f"IMP_{ano - 1}.csv")
    df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv")
    df_exp_mun = carregar_dataframe(url_exp_mun, f"EXP_{ano}_MUN.csv")
    df_imp_mun = carregar_dataframe(url_imp_mun, f"IMP_{ano}_MUN.csv")
    df_uf_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv")

    if any(df is None or df.empty for df in
           [df_exp_ano, df_exp_ano_anterior, df_imp_ano, df_imp_ano_anterior, df_ncm, df_exp_mun, df_imp_mun,
            df_uf_mun]):
        print("Erro ao carregar um ou mais arquivos CSV.")
        return None

    ultimo_mes_disponivel = df_exp_ano['CO_MES'].max()
    ano_completo = ultimo_mes_disponivel == 12

    # --- EXPORTAÇÃO ---
    df_exp_ano_estados = filtrar_dados_por_estado_e_mes(df_exp_ano, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
    df_exp_ano_anterior_estados = filtrar_dados_por_estado_e_mes(df_exp_ano_anterior, estados_brasileiros, ultimo_mes_disponivel, ano_completo)

    df_exp_ano_mg = filtrar_dados_por_estado_e_mes(df_exp_ano, ['MG'], ultimo_mes_disponivel, ano_completo)
    df_exp_ano_anterior_mg = filtrar_dados_por_estado_e_mes(df_exp_ano_anterior, ['MG'], ultimo_mes_disponivel, ano_completo)

    df_exp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)
    df_exp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)

    exportacao_pais_ano = df_exp_ano_mg_paises['VL_FOB'].sum()
    exportacao_pais_ano_anterior = df_exp_ano_anterior_mg_paises['VL_FOB'].sum()

    soma_estados_exp_ano, soma_estados_exp_ano_anterior = calcular_soma_por_estado(df_exp_ano_estados, ano, ano_completo, ultimo_mes_disponivel, df_exp_ano_anterior_estados)

    classificacao_estados_exp = calcular_classificacao_estados(soma_estados_exp_ano, soma_estados_exp_ano_anterior, ano)

    posicao_mg_exp = calcular_posicao_mg(classificacao_estados_exp)

    if not agrupado:
        ranking_mg_exp = calcular_ranking_por_pais(df_exp_ano_mg, ano_completo, ultimo_mes_disponivel)
        try:
            posicao_pais_mg_exp = ranking_mg_exp.index.get_loc(codigos_paises[0]) + 1
        except KeyError:
            posicao_pais_mg_exp = 0
    else:
        posicao_pais_mg_exp = None 

    exportacao_mg_total_ano = 0
    df_exp_ano_mg_filtered = df_exp_ano_mg
    if not ano_completo:
        df_exp_ano_mg_filtered = df_exp_ano_mg[df_exp_ano_mg['CO_MES'] <= ultimo_mes_disponivel]
    exportacao_mg_total_ano = df_exp_ano_mg_filtered['VL_FOB'].sum()

    participacao_pais_mg_exp = calcular_participacao(exportacao_pais_ano, exportacao_mg_total_ano)

    diferenca_exportacao, tipo_diferenca_exp = calcular_diferenca_percentual(exportacao_pais_ano, exportacao_pais_ano_anterior, ano_completo, ultimo_mes_disponivel)

    df_exp_brasil_periodo = df_exp_ano
    if not ano_completo:
        df_exp_brasil_periodo = df_exp_ano[df_exp_ano['CO_MES'] <= ultimo_mes_disponivel]

    exportacao_mg_para_pais = df_exp_brasil_periodo[
        (df_exp_brasil_periodo['SG_UF_NCM'] == 'MG') &
        (df_exp_brasil_periodo['CO_PAIS'].isin(codigos_paises))
    ]['VL_FOB'].sum()

    exportacao_brasil_pais = df_exp_brasil_periodo[
        df_exp_brasil_periodo['CO_PAIS'].isin(codigos_paises)
    ]['VL_FOB'].sum()

    participacao_mg_brasil_exp = calcular_participacao(exportacao_mg_para_pais, exportacao_brasil_pais)

    posicao_mg_pais_exp = calcular_posicao_estado_pais(df_exp_ano_estados, codigos_paises, ano_completo, ultimo_mes_disponivel)

    # --- IMPORTAÇÃO ---
    df_imp_ano_estados = filtrar_dados_por_estado_e_mes(df_imp_ano, estados_brasileiros, ultimo_mes_disponivel, ano_completo)
    df_imp_ano_anterior_estados = filtrar_dados_por_estado_e_mes(df_imp_ano_anterior, estados_brasileiros, ultimo_mes_disponivel, ano_completo)

    df_imp_ano_mg = filtrar_dados_por_estado_e_mes(df_imp_ano, ['MG'], ultimo_mes_disponivel, ano_completo)
    df_imp_ano_anterior_mg = filtrar_dados_por_estado_e_mes(df_imp_ano_anterior, ['MG'], ultimo_mes_disponivel, ano_completo)

    df_imp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)
    df_imp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior, codigos_paises, agrupado, ultimo_mes_disponivel, ano_completo)

    importacao_pais_ano = df_imp_ano_mg_paises['VL_FOB'].sum()
    importacao_pais_ano_anterior = df_imp_ano_anterior_mg_paises['VL_FOB'].sum()

    soma_estados_imp_ano, soma_estados_imp_ano_anterior = calcular_soma_por_estado(df_imp_ano_estados, ano, ano_completo, ultimo_mes_disponivel, df_imp_ano_anterior_estados)

    classificacao_estados_imp = calcular_classificacao_estados(soma_estados_imp_ano, soma_estados_imp_ano_anterior, ano)

    posicao_mg_imp = calcular_posicao_mg(classificacao_estados_imp)

    if not agrupado:
        ranking_mg_imp = calcular_ranking_por_pais(df_imp_ano_mg, ano_completo, ultimo_mes_disponivel)
        try:
            # Tenta obter a posição do país no ranking
            posicao_pais_mg_imp = ranking_mg_imp.index.get_loc(codigos_paises[0]) + 1
        except KeyError:
            # Se o país não estiver no ranking (comércio zero), define a posição como 0
            posicao_pais_mg_imp = 0
    else:
        posicao_pais_mg_imp = None

    importacao_mg_total_ano = 0
    df_imp_ano_mg_filtered = df_imp_ano_mg
    if not ano_completo:
        df_imp_ano_mg_filtered = df_imp_ano_mg[df_imp_ano_mg['CO_MES'] <= ultimo_mes_disponivel]
    importacao_mg_total_ano = df_imp_ano_mg_filtered['VL_FOB'].sum()
    participacao_pais_mg_imp = calcular_participacao(importacao_pais_ano, importacao_mg_total_ano)

    diferenca_importacao, tipo_diferenca_imp = calcular_diferenca_percentual(importacao_pais_ano, importacao_pais_ano_anterior, ano_completo, ultimo_mes_disponivel)

    df_imp_brasil_periodo = df_imp_ano
    if not ano_completo:
        df_imp_brasil_periodo = df_imp_ano[df_imp_ano['CO_MES'] <= ultimo_mes_disponivel]

    importacao_mg_para_pais = df_imp_brasil_periodo[
        (df_imp_brasil_periodo['SG_UF_NCM'] == 'MG') &
        (df_imp_brasil_periodo['CO_PAIS'].isin(codigos_paises))
    ]['VL_FOB'].sum()

    importacao_brasil_pais = df_imp_brasil_periodo[
        df_imp_brasil_periodo['CO_PAIS'].isin(codigos_paises)
    ]['VL_FOB'].sum()

    participacao_mg_brasil_imp = calcular_participacao(importacao_mg_para_pais, importacao_brasil_pais)

    posicao_mg_pais_imp = calcular_posicao_estado_pais(df_imp_ano_estados, codigos_paises, ano_completo, ultimo_mes_disponivel)

    # --- GERAL ---
    balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo = calcular_balanca_e_fluxo(exportacao_pais_ano, importacao_pais_ano, exportacao_pais_ano_anterior, importacao_pais_ano_anterior)

    # --- MUNICIPIOS ---
    df_exp_mun, df_imp_mun = filtrar_dados_municipios(df_exp_mun, df_imp_mun, codigos_paises, ultimo_mes_disponivel, ano_completo)

    exportacoes_por_municipio, total_exportacoes_municipios = agregar_dados_por_municipio(df_exp_mun, ano_completo, ultimo_mes_disponivel)
    importacoes_por_municipio, total_importacoes_municipios = agregar_dados_por_municipio(df_imp_mun, ano_completo, ultimo_mes_disponivel)

    # --- PRODUTOS ---
    produtos_exportacao = agregar_dados_por_produto(df_exp_ano_mg_paises.copy(), df_ncm, ano_completo, ultimo_mes_disponivel) # Usado .copy() para evitar o SettingWithCopyWarning ao modificar o DataFrame na função.
    produtos_importacao = agregar_dados_por_produto(df_imp_ano_mg_paises.copy(), df_ncm, ano_completo, ultimo_mes_disponivel) # Usado .copy() para evitar o SettingWithCopyWarning ao modificar o DataFrame na função.


    return (balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo,
            nomes_paises_validos, posicao_mg_pais_exp, participacao_mg_brasil_exp, exportacao_pais_ano, exportacao_pais_ano_anterior,
            tipo_diferenca_exp, diferenca_exportacao, posicao_pais_mg_exp, participacao_pais_mg_exp, posicao_mg_pais_imp,
            participacao_mg_brasil_imp, importacao_pais_ano, importacao_pais_ano_anterior, tipo_diferenca_imp, diferenca_importacao,
            posicao_pais_mg_imp, participacao_pais_mg_imp, ano_completo, ultimo_mes_disponivel, produtos_exportacao, exportacao_pais_ano, # Changed produtos_exportacao_total to exportacao_pais_ano
            produtos_importacao, importacao_pais_ano, df_ncm, exportacoes_por_municipio, total_exportacoes_municipios, # Changed produtos_importacao_total to importacao_pais_ano
            importacoes_por_municipio, total_importacoes_municipios, df_uf_mun)


def obter_artigo_pais_gemini(nome_pais):
    print(f"Consultando IA para obter o artigo de '{nome_pais}'...")
    prompt = f"Qual o artigo para me referir ao território \"{nome_pais}\". Responda somente com o artigo"

    url = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash-lite:generateContent?key=' + api_key
    headers = {'Content-Type': 'application/json'}
    data = {'contents': [{'parts': [{'text': prompt}]}]}

    try:
        resposta = requests.post(url, headers=headers, json=data, timeout=20)
        resposta.raise_for_status()

        conteudo_resposta = resposta.json()
        if 'candidates' in conteudo_resposta and conteudo_resposta['candidates']:
            texto_bruto = conteudo_resposta['candidates'][0]['content']['parts'][0]['text']
            artigo = texto_bruto.strip().replace('.', '').lower()
            return artigo
        else:
            print("A API não retornou um candidato válido para o artigo.")
            return None
    except requests.exceptions.RequestException as e:
        print(f"Erro na chamada da API para obter artigo: {e}")
        return None


def chamar_gemini(texto):
    """Chama a API do Google Gemini para processar o texto."""
    url = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash-lite:generateContent?key=' + api_key
    headers = {'Content-Type': 'application/json'}
    data = {
        'contents': [
            {
                'parts': [
                    {'text': texto}
                ]
            }
        ]
    }

    try:
        resposta = requests.post(url, headers=headers, json=data, timeout=60)
        resposta.raise_for_status()

        conteudo_resposta = resposta.json()
        if conteudo_resposta.get('candidates'):
            texto_processado = conteudo_resposta['candidates'][0]['content']['parts'][0]['text']
            paragraphs = texto_processado.split('\n\n')
            return [p.strip() for p in paragraphs if p.strip()]
        else:
            print("A API não retornou nenhuma informação. Verifique a entrada e os parâmetros da API.")
            return None
    except requests.exceptions.RequestException as e:
        print(f"Erro na chamada da API principal: {e}")
        return None

def formatar_valor(valor):
    prefixo = ""
    if valor < 0:
        prefixo = "-"
        valor = abs(valor)

    # Bilhões
    if valor >= 999_500_000:
        valor_em_bilhoes = round(valor / 1_000_000_000, 2)
        valor_formatado_str = f"{valor_em_bilhoes:,.2f}"

        unidade = "bilhão" if valor_em_bilhoes < 2 else "bilhões"
        resultado = f"US$ {valor_formatado_str} {unidade}"

    # Milhões
    elif valor >= 999_500:
        valor_em_milhoes = round(valor / 1_000_000, 2)
        valor_formatado_str = f"{valor_em_milhoes:,.2f}"

        unidade = "milhão" if valor_em_milhoes < 2 else "milhões"
        resultado = f"US$ {valor_formatado_str} {unidade}"

    # Mil
    elif valor >= 1_000:
        valor_formatado_str = f"{valor / 1_000:,.1f}"
        resultado = f"US$ {valor_formatado_str} mil"

    # Menor que 1.000
    else:
        resultado = f"US$ {valor:.0f}"

    resultado = resultado.replace(',', 'X').replace('.', ',').replace('X', '.')

    return f"{prefixo}{resultado}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

class DocumentoApp:
    def __init__(self):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""

    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_conteudo_formatado(self, texto):
        """Adiciona um parágrafo de texto formatado ao documento."""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_conteudo_central(self, texto):
        """Adiciona um parágrafo de texto centralizado ao documento."""
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_paragrafo(self, texto):
        """Adiciona um novo parágrafo ao documento."""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_titulo(self, texto):
        """Adiciona um título de seção ou subseção."""
        p = self.doc.add_paragraph()
        if self.subsecao_atual == 0:
            # É uma seção
            run = p.add_run(f"{self.secao_atual}. {texto}")
        else:
            # É uma subsecao
            run = p.add_run(f"{self.secao_atual}.{self.subsecao_atual}. {texto}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def nova_secao(self):
        """Inicia uma nova seção."""
        self.secao_atual += 1
        self.subsecao_atual = 0

    def nova_subsecao(self):
        """Inicia uma nova subseção."""
        self.subsecao_atual += 1

    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)

        header = section.header

        largura_total_cm = 15.88
        table = header.add_table(rows=1, cols=2, width=Cm(largura_total_cm))

        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.columns[0].width = Cm(2.91)
        table.columns[1].width = Cm(12.97)

        cell_imagem = table.cell(0, 0)

        paragraph_imagem = cell_imagem.paragraphs[0]
        paragraph_imagem.paragraph_format.space_before = Pt(0)
        paragraph_imagem.paragraph_format.space_after = Pt(0)

        run_imagem = paragraph_imagem.add_run()
        run_imagem.add_picture('/content/drive/MyDrive/Comex/LogoMinasGerais.png',
                               width=Inches(1.71 / 2.54),
                               height=Inches(1.67 / 2.54))
        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]

        p = cell_texto.paragraphs[0]
        p.text = textos[0]

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = True

        for texto in textos[1:]:
            p = cell_texto.add_paragraph(texto)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

    def finalizar_documento(self):
        os.makedirs(diretorio, exist_ok=True)

        nome_arquivo = f"{self.titulo_doc}.docx"
        nome_arquivo_sanitizado = sanitize_filename(nome_arquivo)
        caminho_completo = os.path.join(diretorio, nome_arquivo_sanitizado)

        self.doc.save(caminho_completo)
        print(f"Documento salvo em: {caminho_completo}")


def main():

    app = DocumentoApp()

    while True:
        ano = int(input("Digite o ano: "))
        if ano < 1998:
            print("Dados não disponíveis para anos anteriores à 1998. Favor inserir novo ano:")
            continue

        url_exp_ano = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano}.csv"
        resposta_exp = requests.head(url_exp_ano, verify=False, timeout=10)
        if resposta_exp.status_code != 200:
            print("Não foram encontrados arquivos de dados para o ano inserido. Favor inserir novo ano:")
            continue

        paises_input = input("Digite os países (separados por ponto e vírgula e espaço): ")
        paises = [pais.strip() for pais in paises_input.split('; ')]

        if len(paises) > 1:
            while True:
                agrupamento_input = input("Deseja que os dados sejam agrupados ou separados? (agrupados/separados): ").lower()
                if agrupamento_input in ("agrupados", "separados"):
                    agrupado = agrupamento_input == "agrupados"
                    break
                else:
                    print("Opção inválida. Digite 'agrupados' ou 'separados'.")
        else:
            agrupado = False  

        if agrupado:
            while True:
                nome_agrupamento = input("Deseja dar um nome para este agrupamento de países? (s/n): ").lower()
                if nome_agrupamento == 's':
                    nome_agrupamento = input("Digite o nome do agrupamento: ")
                    break
                elif nome_agrupamento == 'n':
                    nome_agrupamento = None  
                    break
                else:
                    print("Opção inválida. Digite 's' ou 'n'.")

            resultados = calcular_fluxo_comercial(ano, paises, agrupado=True)
            if resultados is not None:
                (balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca,
                 variacao_fluxo, paises_corretos, posicao_mg_pais_exp, participacao_mg_brasil_exp, exportacao_pais_ano,
                 exportacao_pais_ano_anterior, tipo_diferenca_exp, diferenca_exportacao, posicao_pais_mg_exp,
                 participacao_pais_mg_exp, posicao_mg_pais_imp, participacao_mg_brasil_imp, importacao_pais_ano,
                 importacao_pais_ano_anterior, tipo_diferenca_imp, diferenca_importacao, posicao_pais_mg_imp,
                 participacao_pais_mg_imp, ano_completo, ultimo_mes_disponivel, produtos_exportacao, total_produtos_exp,
                 produtos_importacao, total_produtos_imp, df_ncm, exportacoes_por_municipio,
                 total_exportacoes_municipios, importacoes_por_municipio, total_importacoes_municipios, df_uf_mun) = resultados

                nome_relatorio = nome_agrupamento if nome_agrupamento else ', '.join(paises_corretos)

                if ano_completo:
                    fluxo_e_balanca = f"Em {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, representando {'aumento' if variacao_fluxo > 0 else 'queda'} de {abs(variacao_fluxo):.2f}% em comparação a {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação a {ano-1}."
                    frase_1 = f"Em {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, representando {'aumento' if variacao_fluxo > 0 else 'queda'} de {abs(variacao_fluxo):.2f}% em comparação a {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação a {ano-1}."
                else:
                    fluxo_e_balanca = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, com {'um aumento' if variacao_fluxo > 0 else 'uma queda'} de {abs(variacao_fluxo):.2f}% em comparação ao mesmo período em {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação ao mesmo período em {ano-1}."
                    frase_1 = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, com {'um aumento' if variacao_fluxo > 0 else 'uma queda'} de {abs(variacao_fluxo):.2f}% em comparação ao mesmo período em {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação ao mesmo período em {ano-1}."

                # --- EXPORTAÇÃO (AGRUPADO) ---
                if ano_completo:
                    texto_exportacao = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano}, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio} nas exportações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_exp}%. "
                    frase_2 = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano}, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio} nas exportações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_exp}%. "
                else:
                    texto_exportacao = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} neste período, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio} nas exportações totais de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano}, foi equivalente a {participacao_pais_mg_exp}%. "
                    frase_2 = f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} neste período, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio} nas exportações totais de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano}, foi equivalente a {participacao_pais_mg_exp}%. "

                if ano_completo:
                    texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao(s) país(es)."
                    frase_3 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao(s) país(es)."
                else:
                    texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao(s) país(es)."
                    frase_3 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao(s) país(es)."

                # Produtos Exportação (AGRUPADO)
                if ano_completo:
                    texto_produtos_exportacao = f"Em {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                    frase_6 = f"Em {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                else:
                    texto_produtos_exportacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                    frase_6 = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "

                texto_produtos_exportacao_lista = []
                frase_6_lista = []
                for nome_produto, valor_fob in produtos_exportacao.items(): 
                    participacao_produto_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) 
                    texto_produtos_exportacao_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")
                    frase_6_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")

                texto_produtos_exportacao += "; ".join(texto_produtos_exportacao_lista) + "."
                frase_6 += "; ".join(frase_6_lista) + "."


                # Municípios exportação (AGRUPADO)
                if ano_completo:
                    texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano}, os principais foram: "
                    frase_7 = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano}, os principais foram: "
                else:
                    texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "
                    frase_7 = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "

                texto_municipios_exportacao_lista = []
                frase_7_lista = []
                for i, (codigo_municipio, valor_fob) in enumerate(exportacoes_por_municipio.head(5).items()):
                    nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                    participacao_municipio_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) # Cálculo correto da participação
                    texto_municipios_exportacao_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")
                    frase_7_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")

                texto_municipios_exportacao += "; ".join(texto_municipios_exportacao_lista) + "."
                frase_7 += "; ".join(frase_7_lista) + "."


                # --- IMPORTAÇÃO (AGRUPADO) ---
                if ano_completo:
                    texto_importacao = f"As importações mineiras provenientes {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} em {ano}, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                    frase_4 = f"As importações mineiras provenientes {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} em {ano}, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                else:
                    texto_importacao = f"As importações mineiras provenientes {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} neste período, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                    frase_4 = f"As importações mineiras provenientes {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} neste período, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "

                if ano_completo:
                    texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio} em {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao(s) país(es)."
                    frase_5 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio} em {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao(s) país(es)."
                else:
                    texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao(s) país(es)."
                    frase_5 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao(s) país(es)."

                # Produtos importação (AGRUPADO)
                if ano_completo:
                    texto_produtos_importacao = f"Em {ano}, os principais produtos importados para Minas Gerais {nome_relatorio} foram: "
                    frase_8 = f"Em {ano}, os principais produtos importados para Minas Gerais {nome_relatorio} foram: "
                else:
                    texto_produtos_importacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos importados para Minas Gerais {nome_relatorio} foram: "
                    frase_8 = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos importados para Minas Gerais {nome_relatorio} foram: "

                texto_produtos_importacao_lista = []
                frase_8_lista = []
                for nome_produto, valor_fob in produtos_importacao.items():
                    participacao_produto_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                    texto_produtos_importacao_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")
                    frase_8_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")

                texto_produtos_importacao += "; ".join(texto_produtos_importacao_lista) + "."
                frase_8 += "; ".join(frase_8_lista) + "."

                # Municípios importação (AGRUPADO)
                if ano_completo:
                    texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio} em {ano}, os principais foram: "
                    frase_9 = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio} em {ano}, os principais foram: "
                else:
                    texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "
                    frase_9 = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "

                texto_municipios_importacao_lista = []
                frase_9_lista = []
                for i, (codigo_municipio, valor_fob) in enumerate(importacoes_por_municipio.head(5).items()):
                    nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                    participacao_municipio_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                    texto_municipios_importacao_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")
                    frase_9_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")

                texto_municipios_importacao += "; ".join(texto_municipios_importacao_lista) + "."
                frase_9 += "; ".join(frase_9_lista) + "."


                # --- CORREÇÃO DE TEXTO GEMINI (AGRUPADO) ---
                texto_relatorio = frase_1 + '\n' + frase_2 + frase_3 + frase_6 + frase_7 + '\n' + frase_4 + frase_5 + frase_8 + frase_9

                texto_processado_ia_paragraphs = chamar_gemini(f"Agrupe todos os pontos em 5 parágrafos, relacionando por assunto. Você não pode suprimir nenhuma das informações e não pode adicionar nenhuma palavra ou texto que forneça qualquer tipo de valoração ou juízo de valor. Ou seja, sua função é apenas transformar o texto de tópicos para parágrafos. A seguir, o texto: \n{texto_relatorio}")

                if nome_agrupamento:
                    titulo_documento = f"Briefing - {nome_agrupamento} - {ano}"
                else:
                    titulo_documento = f"Briefing - {paises_corretos[0]} e outro(s) {len(paises_corretos)-1} país(es) - {ano}"

                app.set_titulo(titulo_documento)
                app.nova_secao()
                app.adicionar_titulo("Fluxo Comercial")
                app.adicionar_conteudo_formatado(fluxo_e_balanca)
                app.nova_secao()
                app.adicionar_titulo("Exportações")
                app.adicionar_conteudo_formatado(texto_exportacao)
                app.adicionar_conteudo_formatado(texto_exportacao_2)
                app.adicionar_conteudo_formatado(texto_produtos_exportacao)
                app.adicionar_conteudo_formatado(texto_municipios_exportacao)
                app.nova_secao()
                app.adicionar_titulo("Importações")
                app.adicionar_conteudo_formatado(texto_importacao)
                app.adicionar_conteudo_formatado(texto_importacao_2)
                app.adicionar_conteudo_formatado(texto_produtos_importacao)
                app.adicionar_conteudo_formatado(texto_municipios_importacao)
                app.nova_secao()
                app.adicionar_titulo("Texto processado pela IA")
                if texto_processado_ia_paragraphs: 
                    for paragraph in texto_processado_ia_paragraphs:
                        app.adicionar_conteudo_formatado(paragraph)
                app.finalizar_documento()


                print("-" * 30)  
                break  

        else:
            for pais in paises:
                app = DocumentoApp()

                resultados = calcular_fluxo_comercial(ano, [pais], agrupado=False)

                if resultados is not None:
                    (balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior,
                     variacao_balanca, variacao_fluxo, pais_correto, posicao_mg_pais_exp, participacao_mg_brasil_exp,
                     exportacao_pais_ano, exportacao_pais_ano_anterior, tipo_diferenca_exp, diferenca_exportacao,
                     posicao_pais_mg_exp, participacao_pais_mg_exp, posicao_mg_pais_imp, participacao_mg_brasil_imp,
                     importacao_pais_ano, importacao_pais_ano_anterior, tipo_diferenca_imp, diferenca_importacao,
                     posicao_pais_mg_imp, participacao_pais_mg_imp, ano_completo, ultimo_mes_disponivel,
                     produtos_exportacao, total_produtos_exp, produtos_importacao, total_produtos_imp,
                     df_ncm, exportacoes_por_municipio, total_exportacoes_municipios, importacoes_por_municipio,
                     total_importacoes_municipios, df_uf_mun) = resultados

                    nome_pais_base = pais_correto[0]

                    # --- ARTIGO ---
                    nome_relatorio = nome_pais_base
                    nome_relatorio_capitalizado = nome_pais_base
                    artigo = obter_artigo_pais_gemini(nome_pais_base)
                    valid_articles = ['o', 'a', 'os', 'as']

                    if artigo and artigo.lower() in valid_articles:
                        nome_relatorio = f"{artigo.lower()} {nome_pais_base}"
                        nome_relatorio_capitalizado = f"{artigo.capitalize()} {nome_pais_base}"
                    else:
                        artigo = None
                        print(f"Não foi possível obter um artigo válido para '{nome_pais_base}'. Usando nome do país sem artigo.")

                    contracoes_map = {
                        'o': 'do',
                        'a': 'da',
                        'os': 'dos',
                        'as': 'das'
                    }

                    preposicao_contraida = contracoes_map.get(artigo)

                    if preposicao_contraida:
                        nome_relatorio_com_contracao = f"{preposicao_contraida} {nome_pais_base}"
                    else:
                        nome_relatorio_com_contracao = f"de {nome_pais_base}"

                    titulo_documento = f"Briefing - {nome_pais_base} - {ano}"
                    app.set_titulo(titulo_documento)


                    if ano_completo:
                        fluxo_e_balanca = f"Em {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, representando {'aumento' if variacao_fluxo > 0 else 'queda'} de {abs(variacao_fluxo):.2f}% em comparação a {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação a {ano-1}."
                        frase_1 = f"Em {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, representando {'aumento' if variacao_fluxo > 0 else 'queda'} de {abs(variacao_fluxo):.2f}% em comparação a {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação a {ano-1}."
                    else:
                        fluxo_e_balanca = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, com {'um aumento' if variacao_fluxo > 0 else 'uma queda'} de {abs(variacao_fluxo):.2f}% em comparação ao mesmo período em {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação ao mesmo período em {ano-1}."
                        frase_1 = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, com {'um aumento' if variacao_fluxo > 0 else 'uma queda'} de {abs(variacao_fluxo):.2f}% em comparação ao mesmo período em {ano-1}. A balança comercial fechou {'positiva' if balanca_ano > 0 else 'negativa'} para Minas Gerais em {formatar_valor(balanca_ano)}, apresentando {'um crescimento' if variacao_balanca > 0 else 'uma queda'} de {abs(variacao_balanca):.1f}% em relação ao mesmo período em {ano-1}."

                    # --- EXPORTAÇÃO (SEM AGRUPAMENTO) ---
                    if posicao_pais_mg_exp > 0: # País está no ranking
                        if ano_completo:
                            texto_exportacao = f"{nome_relatorio_capitalizado} foi o {posicao_pais_mg_exp}º destino das exportações de Minas Gerais em {ano}. As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano}, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio_com_contracao} nas exportações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_exp}%. "
                            frase_2 = f"{nome_relatorio_capitalizado} foi o {posicao_pais_mg_exp}º destino das exportações de Minas Gerais em {ano}. As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano}, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio_com_contracao} nas exportações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_exp}%. "
                        else:
                            texto_exportacao = f"{nome_relatorio_capitalizado} foi o {posicao_pais_mg_exp}º destino das exportações de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano}. As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} neste período, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio_com_contracao} nas exportações totais de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano}, foi equivalente a {participacao_pais_mg_exp}%. "
                            frase_2 = f"{nome_relatorio_capitalizado} foi o {posicao_pais_mg_exp}º destino das exportações de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano}. As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} neste período, {tipo_diferenca_exp} de {diferenca_exportacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio_com_contracao} nas exportações totais de Minas Gerais, considerando até {meses_pt[ultimo_mes_disponivel]} de {ano}, foi equivalente a {participacao_pais_mg_exp}%. "
                    else: 
                        if ano_completo:
                            texto_exportacao = f"Em {ano}, Minas Gerais não registrou exportações para {nome_relatorio}."
                            frase_2 = texto_exportacao
                        else:
                            texto_exportacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, Minas Gerais não registrou exportações para {nome_relatorio}."
                            frase_2 = texto_exportacao

                    if posicao_mg_pais_exp > 0: 
                        if ano_completo:
                            texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao país."
                            frase_3 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao país."
                        else:
                            texto_exportacao_2 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao país."
                            frase_3 = f"Minas Gerais foi o {posicao_mg_pais_exp}º principal estado exportador brasileiro para {nome_relatorio} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_exp}% nas vendas do Brasil ao país."
                    else: # Se MG não teve exportações para o país (participação é 0%)
                        texto_exportacao_2 = f"O estado de Minas Gerais não se posicionou no ranking de exportadores brasileiros para {nome_relatorio}, pois não houve registro de vendas."
                        frase_3 = texto_exportacao_2

                    # Produtos Exportação (SEM AGRUPAMENTO)
                    if ano_completo:
                        texto_produtos_exportacao = f"Em {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                        frase_6 = f"Em {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                    else:
                        texto_produtos_exportacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "
                        frase_6 = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: "

                    texto_produtos_exportacao_lista = []
                    frase_6_lista = []
                    for nome_produto, valor_fob in produtos_exportacao.items(): 
                        participacao_produto_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) 
                        texto_produtos_exportacao_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")
                        frase_6_lista.append(f"{nome_produto} ({participacao_produto_exportacao}%)")

                    texto_produtos_exportacao += "; ".join(texto_produtos_exportacao_lista) + "."
                    frase_6 += "; ".join(frase_6_lista) + "."


                    # Municípios Exportação (SEM AGRUPAMENTO)
                    if ano_completo:
                        texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano}, os principais foram: "
                        frase_7 = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano}, os principais foram: "
                    else:
                        texto_municipios_exportacao = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "
                        frase_7 = f"Dentre os {len(exportacoes_por_municipio)} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "

                    texto_municipios_exportacao_lista = []
                    frase_7_lista = []
                    for i, (codigo_municipio, valor_fob) in enumerate(exportacoes_por_municipio.head(5).items()):
                        nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                        participacao_municipio_exportacao = calcular_participacao(valor_fob, exportacao_pais_ano) 
                        texto_municipios_exportacao_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")
                        frase_7_lista.append(f"{nome_municipio} ({participacao_municipio_exportacao}%)")

                    texto_municipios_exportacao += "; ".join(texto_municipios_exportacao_lista) + "."
                    frase_7 += "; ".join(frase_7_lista) + "."

                    # --- IMPORTAÇÃO (SEM AGRUPAMENTO) ---
                    if posicao_pais_mg_imp > 0: 
                        if ano_completo:
                            texto_importacao = f"{nome_relatorio_capitalizado} foi a {posicao_pais_mg_imp}ª origem das importações de Minas Gerais em {ano}. As importações mineiras provenientes {nome_relatorio_com_contracao} somaram {formatar_valor(importacao_pais_ano)} em {ano}, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio_com_contracao} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                            frase_4 = f"{nome_relatorio_capitalizado} foi a {posicao_pais_mg_imp}ª origem das importações de Minas Gerais em {ano}. As importações mineiras provenientes {nome_relatorio_com_contracao} somaram {formatar_valor(importacao_pais_ano)} em {ano}, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação a {ano-1}. A participação {nome_relatorio_com_contracao} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                        else:
                            texto_importacao = f"{nome_relatorio_capitalizado} foi a {posicao_pais_mg_imp}ª origem das importações de Minas Gerais até {meses_pt[ultimo_mes_disponivel]} de {ano}. As importações mineiras provenientes {nome_relatorio_com_contracao} somaram {formatar_valor(importacao_pais_ano)} neste período, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio_com_contracao} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                            frase_4 = f"{nome_relatorio_capitalizado} foi a {posicao_pais_mg_imp}ª origem das importações de Minas Gerais até {meses_pt[ultimo_mes_disponivel]} de {ano}. As importações mineiras provenientes {nome_relatorio_com_contracao} somaram {formatar_valor(importacao_pais_ano)} neste período, {tipo_diferenca_imp} de {diferenca_importacao:.1f}% em relação ao mesmo período em {ano-1}. A participação {nome_relatorio_com_contracao} nas importações totais de Minas Gerais em {ano} foi equivalente a {participacao_pais_mg_imp}%. "
                    else: 
                        if ano_completo:
                            texto_importacao = f"Em {ano}, Minas Gerais não registrou importações provenientes {nome_relatorio_com_contracao}."
                            frase_4 = texto_importacao
                        else:
                            texto_importacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, Minas Gerais não registrou importações provenientes {nome_relatorio_com_contracao}."
                            frase_4 = texto_importacao

                    if posicao_mg_pais_imp > 0: 
                        if ano_completo:
                            texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio_com_contracao} em {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao país."
                            frase_5 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio_com_contracao} em {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao país."
                        else:
                            texto_importacao_2 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio_com_contracao} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao país."
                            frase_5 = f"Minas Gerais foi o {posicao_mg_pais_imp}º principal estado importador brasileiro {nome_relatorio_com_contracao} no período até {meses_pt[ultimo_mes_disponivel]} de {ano}, com uma participação de {participacao_mg_brasil_imp}% nas compras do Brasil ao país."
                    else: 
                        texto_importacao_2 = f"O estado de Minas Gerais não se posicionou no ranking de importadores brasileiros {nome_relatorio_com_contracao}, pois não houve registro de compras."
                        frase_5 = texto_importacao_2

                    # Produtos Importação (SEM AGRUPAMENTO)
                    if ano_completo:
                        texto_produtos_importacao = f"Em {ano}, os principais produtos importados para Minas Gerais {nome_relatorio_com_contracao} foram: "
                        frase_8 = f"Em {ano}, os principais produtos importados para Minas Gerais {nome_relatorio_com_contracao} foram: "
                    else:
                        texto_produtos_importacao = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos importados para Minas Gerais {nome_relatorio_com_contracao} foram: "
                        frase_8 = f"Até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais produtos importados para Minas Gerais {nome_relatorio_com_contracao} foram: "

                    texto_produtos_importacao_lista = []
                    frase_8_lista = []
                    for nome_produto, valor_fob in produtos_importacao.items():
                        participacao_produto_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                        texto_produtos_importacao_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")
                        frase_8_lista.append(f"{nome_produto} ({participacao_produto_importacao}%)")

                    texto_produtos_importacao += "; ".join(texto_produtos_importacao_lista) + "."
                    frase_8 += "; ".join(frase_8_lista) + "."


                    # Municípios Importação (SEM AGRUPAMENTO)
                    if ano_completo:
                        texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio_com_contracao} em {ano}, os principais foram: "
                        frase_9 = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos para {nome_relatorio} em {ano}, os principais foram: "
                    else:
                        texto_municipios_importacao = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos {nome_relatorio_com_contracao} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "
                        frase_9 = f"Dentre os {len(importacoes_por_municipio)} municípios de Minas Gerais que importaram produtos para {nome_relatorio} até {meses_pt[ultimo_mes_disponivel]} de {ano}, os principais foram: "

                    texto_municipios_importacao_lista = []
                    frase_9_lista = []
                    for i, (codigo_municipio, valor_fob) in enumerate(importacoes_por_municipio.head(5).items()):
                        nome_municipio = df_uf_mun[df_uf_mun['CO_MUN_GEO'] == codigo_municipio]['NO_MUN_MIN'].iloc[0]
                        participacao_municipio_importacao = calcular_participacao(valor_fob, importacao_pais_ano) 
                        texto_municipios_importacao_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")
                        frase_9_lista.append(f"{nome_municipio} ({participacao_municipio_importacao}%)")

                    texto_municipios_importacao += "; ".join(texto_municipios_importacao_lista) + "."
                    frase_9 += "; ".join(frase_9_lista) + "."


                    # --- GEMINI ---

                    texto_processado_ia_paragraphs = None
                    if revisao_texto_gemini:
                      texto_relatorio = frase_1 + '\n' + frase_2 + '\n' + frase_3 + '\n' + frase_6 + '\n' + frase_7 + '\n' + frase_4 + '\n' + frase_5 + '\n' + frase_8 + '\n' + frase_9

                      texto_processado_ia_paragraphs = chamar_gemini(f"Ajuste a ortografia e concordância das orações a seguir. Você não pode suprimir nenhuma das informações e não pode adicionar nenhuma palavra ou texto que forneça qualquer tipo de valoração ou juízo de valor. Ou seja, sua função é apenas fazer ajustes de ortografia e concordância nas orações, mantendo todas as informações. Faça o retorno em formatação simples. A seguir, as orações: \n{texto_relatorio}")

                    else:
                      print("Revisão da IA desativada.")

                    app.nova_secao()
                    app.adicionar_titulo("Fluxo Comercial")
                    app.adicionar_conteudo_formatado(fluxo_e_balanca)

                    app.nova_secao()
                    app.adicionar_titulo("Exportações")
                    app.adicionar_conteudo_formatado(texto_exportacao)
                    app.adicionar_conteudo_formatado(texto_exportacao_2)
                    app.adicionar_conteudo_formatado(texto_produtos_exportacao)
                    app.adicionar_conteudo_formatado(texto_municipios_exportacao)

                    app.nova_secao()
                    app.adicionar_titulo("Importações")
                    app.adicionar_conteudo_formatado(texto_importacao)
                    app.adicionar_conteudo_formatado(texto_importacao_2)
                    app.adicionar_conteudo_formatado(texto_produtos_importacao)
                    app.adicionar_conteudo_formatado(texto_municipios_importacao)

                    if texto_processado_ia_paragraphs:
                        app.nova_secao()
                        app.adicionar_titulo("Texto processado pela IA")
                        for paragraph in texto_processado_ia_paragraphs:
                            app.adicionar_conteudo_formatado(paragraph)

                    app.finalizar_documento()

                    print("-" * 30)  

            break

if __name__ == "__main__":
    main()